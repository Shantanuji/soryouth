import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

heading = doc.add_heading('Perfect Placeholder Blocks', 0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Copy and paste these perfectly aligned tables directly into your main company template. Do not edit the text inside the curly braces {{like_this}}.', style='Normal')

def add_table_block(title, rows_data):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Placeholder'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        shading_elm = parse_xml(r'<w:shd {} w:fill="DCE6F1"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    for field, placeholder in rows_data:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = placeholder
        
    doc.add_paragraph() # spacing

add_table_block('1. Client Details', [
    ('Client Name', '{{name}}'),
    ('Contact Person', '{{contact_person}}'),
    ('Email', '{{email}}'),
    ('Phone', '{{phone}}'),
    ('Location', '{{location}}'),
    ('City/Area', '{{city_area}}'),
    ('Client Type', '{{client_type}}')
])

add_table_block('2. Proposal Information', [
    ('Proposal No', '{{proposal_number}}'),
    ('Date', '{{proposal_date}}'),
    ('Prepared By', '{{created_by}}'),
    ('Today\'s Date', '{{date_today}}')
])

add_table_block('3. System Specifications', [
    ('System Capacity (kW)', '{{capacity}}'),
    ('Module Type', '{{module_type}}'),
    ('Module Wattage', '{{module_wattage}}'),
    ('DCR Status', '{{dcr_status}}'),
    ('Inverter Rating', '{{inverter_rating}}'),
    ('Inverter Qty', '{{inverter_qty}}'),
    ('Space Required (sq.ft)', '{{required_space}}'),
    ('LA Kit Qty', '{{la_kit_qty}}'),
    ('ACDB/DCDB Qty', '{{acdb_dcdb_qty}}'),
    ('Earthing Kit Qty', '{{earthing_kit_qty}}')
])

add_table_block('4. Commercials (Costing)', [
    ('Cost per kW (₹)', '{{cost_per_kw}}'),
    ('Project Cost ex GST (₹)', '{{project_cost_ex_gst}}'),
    ('GST Amount (₹)', '{{gst_amount}}'),
    ('CGST Amount (₹)', '{{cgst_amount}}'),
    ('SGST Amount (₹)', '{{sgst_amount}}'),
    ('Total Cost inc GST (₹)', '{{total_project_cost_inc_gst}}'),
    ('Subsidy Amount (₹)', '{{subsidy_amount}}'),
    ('Net Amount After Subsidy (₹)', '{{net_amount_after_subsidy}}')
])

add_table_block('5. Performance & ROI', [
    ('Annual Generation (kWh)', '{{annual_generation}}'),
    ('Grid Tariff (₹/Unit)', '{{grid_tariff_per_unit}}'),
    ('ROI (Years)', '{{roi_in_years}}')
])

doc.add_heading('6. Dynamic Charts & Sheets', level=2)
doc.add_paragraph('Paste these placeholders exactly where you want the large charts to appear. They will automatically stretch to fit the page horizontally.')
doc.add_paragraph('{{monthly_generation_chart}}')
doc.add_paragraph('{{yearly_savings_chart}}')
doc.add_paragraph('{{capex_evaluation_sheet}}')

doc.save('public/Placeholder_Blocks.docx')
