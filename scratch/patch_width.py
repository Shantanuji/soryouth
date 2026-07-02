import re

with open('microservices/proposal_generator/main.py', 'r', encoding='utf-8') as f:
    content = f.read()

new_helpers = """
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def format_cell(cell, text, bold=False, color_hex='000000', align='left', size=9, bg_color=None):
    if bg_color:
        set_cell_background(cell, bg_color)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

def set_table_width(table, width_in_inches):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(width_in_inches * 1440)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def create_table_section(parent, title, rows, col_widths, title_bg='002060', title_color='FFFFFF'):
    cols = len(col_widths)
    table = parent.add_table(rows=len(rows) + 1, cols=cols)
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass
    table.autofit = False
    set_table_width(table, sum(col_widths))
    
    for i, w in enumerate(col_widths):
        table.columns[i].width = Inches(w)
        for cell in table.columns[i].cells:
            cell.width = Inches(w)
            
    hdr_cells = table.rows[0].cells
    hdr_cells[0].merge(hdr_cells[-1])
    format_cell(hdr_cells[0], title, bold=True, color_hex=title_color, align='center', bg_color=title_bg, size=10)
    
    for i, row_data in enumerate(rows):
        cells = table.rows[i + 1].cells
        is_total = str(row_data[0]).startswith('**')
        label = str(row_data[0]).replace('**', '').strip()
        if is_total:
            bg = '92D050'
        elif i % 2 == 0:
            bg = 'DCE6F1'
        else:
            bg = 'FFFFFF'
            
        format_cell(cells[0], label, bold=is_total, align='left', bg_color=bg)
        for j in range(1, cols):
            val = row_data[j] if j < len(row_data) else ""
            format_cell(cells[j], val, bold=is_total, align='right', bg_color=bg)
    return table

def create_capex_evaluation_sheet(doc, context):
    def sf(key, default=0):
        v = context.get(key, default)
        if v is None:
            return float(default)
        try:
            return float(str(v).replace(',', '').replace('\u20b9', '').replace('%', '').strip())
        except (ValueError, TypeError):
            return float(default)

    name       = str(context.get('name', 'N/A'))
    location   = str(context.get('location', 'N/A'))
    capacity   = sf('capacity')
    rate_pw    = sf('rate_per_watt')
    cost_pkw   = rate_pw * 1000
    base_amt   = sf('base_amount') or (cost_pkw * capacity)
    final_amt  = sf('final_amount') or sf('total_project_cost_inc_gst') or (base_amt * 1.09)
    gst_amt    = max(0.0, final_amt - base_amt)
    gst_pct    = (gst_amt / base_amt * 100) if base_amt > 0 else 9.0
    unit_rate  = sf('unit_rate') or sf('grid_tariff_per_unit')
    subsidy    = sf('subsidy_amount')
    gen_yr     = sf('generation_per_year') or (capacity * 4 * 345)
    savings_yr = sf('savings_per_year') or (gen_yr * unit_rate)

    ad1 = base_amt * 0.40 * 0.25
    ad2 = base_amt * 0.60 * 0.60 * 0.25
    ad3 = base_amt * 0.60 * 0.40 * 0.675 * 0.25
    total_ad = ad1 + ad2 + ad3
    om_pkw  = 750
    om_base = capacity * om_pkw
    net_inv   = final_amt - subsidy - total_ad
    roi_years = net_inv / savings_yr if savings_yr > 0 else 0

    sd = doc.new_subdoc()
    subdoc = sd.subdocx
    
    p = subdoc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Evaluation Sheet for Capex Model")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string('002060')
    
    subdoc.add_paragraph()
    
    master_table = subdoc.add_table(rows=1, cols=3)
    master_table.autofit = False
    
    widths = [2.9, 0.2, 2.9]
    set_table_width(master_table, 6.0)
    for i, w in enumerate(widths):
        master_table.columns[i].width = Inches(w)
        for cell in master_table.columns[i].cells:
            cell.width = Inches(w)
    
    left_cell = master_table.cell(0, 0)
    right_cell = master_table.cell(0, 2)
    
    create_table_section(
        left_cell, "Project Specification", 
        [
            ("Client Name", name, ""),
            ("Project Location", location, ""),
            ("Project Size (KW)", "", f"{capacity:.2f}"),
            ("Cost per KW ex GST", "₹", f"{cost_pkw:,.0f}"),
            ("Project Cost ex GST", "₹", f"{base_amt:,.2f}"),
            (f"GST @ {gst_pct:.1f}%", "₹", f"{gst_amt:,.2f}"),
            ("**Total Project Cost inc GST", "₹", f"{final_amt:,.2f}"),
            ("Client's Current Grid Tariff /Unit", "₹", f"{unit_rate:.2f}"),
        ],
        [1.6, 0.3, 1.0]
    )
    left_cell.add_paragraph()
    
    create_table_section(
        left_cell, "Plant Performance", 
        [
            ("Average Annual Generation (KW)", "", f"{gen_yr:,.2f}"),
            ("Average Monthly Generation (KW)", "", f"{gen_yr / 12:,.2f}"),
            ("Estimated Year-on-Year degradation", "", "0.70 - 0.80%"),
        ],
        [2.0, 0.2, 0.7]
    )
    left_cell.add_paragraph()
    
    create_table_section(
        left_cell, "O&M Cost", 
        [
            ("Cost per KW of maintenance", "₹", f"{om_pkw:,.2f}"),
            ("Total Cost per year", "₹", f"{om_base:,.2f}"),
            ("Yearly escalation in O&M Cost", "", "3.00%"),
        ],
        [1.6, 0.3, 1.0]
    )
    
    create_table_section(
        right_cell, "Accelerated Depreciation Benefits", 
        [
            ("1st year - 25% tax savings on 40% depreciation", "₹", f"{ad1:,.2f}"),
            ("2nd year - 25% tax savings on 40% depreciation", "₹", f"{ad2:,.2f}"),
            ("3rd year - 25% tax savings on 20% depreciation", "₹", f"{ad3:,.2f}"),
            ("**Total", "₹", f"{total_ad:,.2f}"),
        ],
        [1.9, 0.2, 0.8]
    )
    right_cell.add_paragraph()
    
    create_table_section(
        right_cell, "Return On Investment (ROI) Calculation", 
        [
            ("Project Cost ex GST", "₹", f"{base_amt:,.2f}"),
            ("Additional / Subsidy Benefits", "₹", f"{subsidy:,.2f}"),
            ("Cost Via Grid", "₹", f"{savings_yr * 25:,.2f}"),
            ("ROI in Years", "", f"{roi_years:.2f}"),
            ("Monthly Payments", "₹", "-"),
            ("**Total Plant cost inc Interest", "₹", f"{net_inv:,.2f}"),
        ],
        [1.6, 0.3, 1.0]
    )
    
    subdoc.add_paragraph()
    
    cols = ['Period', 'Unit\\nGeneration', 'Cost via\\nGrid', 'Solar Plant\\nEMIs', 'O&M Cost', 'GSC\\nCharges', 'AD Benefit', 'Savings via\\nSolar']
    widths = [0.6, 0.75, 0.9, 0.7, 0.7, 0.75, 0.7, 0.9]
    proj_table = subdoc.add_table(rows=27, cols=8)
    try:
        proj_table.style = 'Table Grid'
    except KeyError:
        pass
    proj_table.autofit = False
    set_table_width(proj_table, 6.0)
    
    for i, w in enumerate(widths):
        proj_table.columns[i].width = Inches(w)
        for cell in proj_table.columns[i].cells:
            cell.width = Inches(w)
            
    for i, text in enumerate(cols):
        format_cell(proj_table.rows[0].cells[i], text, bold=True, color_hex='FFFFFF', align='center', bg_color='002060', size=7.5)
        
    yr_data = []
    for y in range(1, 26):
        gen  = gen_yr * (0.994 ** (y - 1))
        grid = gen * unit_rate
        om   = 0.0 if y == 1 else om_base * (1.03 ** (y - 2))
        ad   = [ad1, ad2, ad3][y - 1] if y <= 3 else 0.0
        sav  = grid + ad - om
        yr_data.append((y, gen, grid, om, ad, sav))
        
    for r in range(1, 26):
        bg = 'DCE6F1' if (r-1) % 2 == 0 else 'FFFFFF'
        cells = proj_table.rows[r].cells
        y, gen, grid, om, ad, sav = yr_data[r-1]
        
        format_cell(cells[0], f"Year {y}", align='center', bg_color=bg, size=7)
        format_cell(cells[1], f"{gen:,.2f}", align='right', bg_color=bg, size=7)
        format_cell(cells[2], f"₹ {grid:,.2f}", align='right', bg_color=bg, size=7)
        format_cell(cells[3], "₹ -", align='right', bg_color=bg, size=7)
        format_cell(cells[4], f"₹ {om:,.2f}" if om > 0 else "₹ -", align='right', bg_color=bg, size=7)
        format_cell(cells[5], "₹ -", align='right', bg_color=bg, size=7)
        format_cell(cells[6], f"₹ {ad:,.2f}" if ad > 0 else "₹ -", align='right', bg_color=bg, size=7)
        format_cell(cells[7], f"₹ {sav:,.2f}", align='right', bg_color=bg, size=7)
        
    cells = proj_table.rows[26].cells
    format_cell(cells[0], "Total over 25 years", bold=True, color_hex='FFFFFF', align='center', bg_color='002060', size=7)
    format_cell(cells[1], f"{sum(r[1] for r in yr_data):,.3f}", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)
    format_cell(cells[2], f"₹ {sum(r[2] for r in yr_data):,.2f}", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)
    format_cell(cells[3], "₹ -", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)
    format_cell(cells[4], f"₹ {sum(r[3] for r in yr_data):,.2f}", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)
    format_cell(cells[5], "₹ -", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)
    format_cell(cells[6], f"₹ {sum(r[4] for r in yr_data):,.2f}", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)
    format_cell(cells[7], f"₹ {sum(r[5] for r in yr_data):,.2f}", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)

    note_text = (
        "The estimates provided in the above table are subject to the following conditions:\\n"
        "• The generation numbers are estimated as per current weather data and will change as per actual weather conditions.\\n"
        "• A standard generation reduction of 0.60% year-on-year is applied for further accuracy.\\n"
        "• The savings shown are based on continuous consumption of the generated power by the client.\\n"
        "• Any change in government net metering policy or other laws may affect the savings and are out of our control.\\n"
        "• Any O&M cost incurred during the plant life will be extra which is calculated in the table.\\n"
        "• Furthermore any unexpected damage due to natural disaster is not factored in the calculations."
    )
    p_note = subdoc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_note.add_run(note_text)
    run.font.name = 'Calibri'
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor.from_string('444444')
    
    return sd
"""

start_idx = -1
end_idx = -1
lines = content.split('\n')
for i, line in enumerate(lines):
    if line.startswith('from docx.shared import Inches'):
        start_idx = i
    if line.startswith('@app.route'):
        if start_idx != -1 and i > start_idx:
            end_idx = i
            break

if start_idx != -1 and end_idx != -1:
    new_content = '\n'.join(lines[:start_idx]) + '\n' + new_helpers + '\n\n' + '\n'.join(lines[end_idx:])
    with open('microservices/proposal_generator/main.py', 'w', encoding='utf-8') as f:
        f.write(new_content)
    print('Updated fully with set_table_width!')
else:
    print('Indices not found!')
