import docx

doc = docx.Document("scratch/template_current.docx")

patched_count = 0
for idx, p in enumerate(doc.paragraphs):
    p_xml = p._p.xml
    if 'LINK Excel' in p_xml and 'Commercial Offer' in p_xml:
        print(f"Patching Paragraph {idx}...")
        # Clear paragraph text and runs
        p.text = ""
        # Add the placeholder run
        p.add_run("{{commercial_offer}}")
        patched_count += 1

print(f"Patched {patched_count} paragraphs.")
if patched_count > 0:
    doc.save("scratch/template_patched.docx")
    print("Saved to scratch/template_patched.docx")
