import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Set background color of a cell using hex string like '002060'"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=0, bottom=0, left=50, right=50):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def format_cell(cell, text, bold=False, color_hex='000000', align='left', size=9, bg_color=None):
    if bg_color:
        set_cell_background(cell, bg_color)
    
    cell.text = ""
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    
    # Vertically center text in cell
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

def create_table_section(parent, title, rows, col_widths, title_bg='002060', title_color='FFFFFF'):
    """rows is a list of tuples: (label, val1, val2...) or just (label, val)"""
    cols = len(col_widths)
    # Using 'Table Grid' gives default black borders
    table = parent.add_table(rows=len(rows) + 1, cols=cols)
    table.style = 'Table Grid'
    
    # Set column widths
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(w)
            
    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].merge(hdr_cells[-1]) # merge all columns for title
    format_cell(hdr_cells[0], title, bold=True, color_hex=title_color, align='center', bg_color=title_bg, size=10)
    
    # Data Rows
    for i, row_data in enumerate(rows):
        cells = table.rows[i + 1].cells
        
        is_total = str(row_data[0]).startswith('**')
        label = str(row_data[0]).replace('**', '').strip()
        
        # Determine background color
        if is_total:
            bg = '92D050' # Green
        elif i % 2 == 0:
            bg = 'DCE6F1' # Light blue
        else:
            bg = 'FFFFFF' # White
            
        format_cell(cells[0], label, bold=is_total, align='left', bg_color=bg)
        
        # Fill rest of columns
        for j in range(1, cols):
            val = row_data[j] if j < len(row_data) else ""
            format_cell(cells[j], val, bold=is_total, align='right', bg_color=bg)
            
    return table

def generate_evaluation_sheet(output_path):
    doc = docx.Document()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Evaluation Sheet for Capex Model")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string('002060')
    
    doc.add_paragraph() # spacing
    
    # Create master layout table (1 row, 3 cols for Left | space | Right)
    master_table = doc.add_table(rows=1, cols=3)
    master_table.columns[0].width = Inches(3.2)
    master_table.columns[1].width = Inches(0.2) # gap
    master_table.columns[2].width = Inches(3.2)
    
    left_cell = master_table.cell(0, 0)
    right_cell = master_table.cell(0, 2)
    
    # Left Side
    create_table_section(
        left_cell, "Project Specification", 
        [
            ("Client Name", "Reddy's", ""),
            ("Project Location", "Rohinjan Gaon", ""),
            ("Project Size (KW)", "", "8.40"),
            ("Cost per KW ex GST", "₹", "62,000"),
            ("Project Cost ex GST", "₹", "5,20,800"),
            ("GST @ 8.9%", "₹", "46,351"),
            ("**Total Project Cost inc GST", "₹", "5,67,151"),
            ("Client's Current Grid Tariff /Unit", "₹", "18.00"),
        ],
        [1.8, 0.4, 1.0]
    )
    left_cell.add_paragraph() # space
    
    create_table_section(
        left_cell, "Plant Performance", 
        [
            ("Average Annual Generation (KW)", "", "11,592.00"),
            ("Average Monthly Generation (KW)", "", "966.00"),
            ("Estimated Year-on-Year degradation", "", "0.70 - 0.80%"),
        ],
        [2.2, 0.2, 0.8]
    )
    left_cell.add_paragraph()
    
    create_table_section(
        left_cell, "O&M Cost", 
        [
            ("Cost per KW of maintenance", "₹", "750.00"),
            ("Total Cost per year", "₹", "6,300.00"),
            ("Yearly escalation in O&M Cost", "", "3.00%"),
        ],
        [1.8, 0.4, 1.0]
    )
    
    # Right Side
    create_table_section(
        right_cell, "Accelerated Depreciation Benefits", 
        [
            ("1st year - 25% tax savings on 40% depreciation", "₹", "52,080.00"),
            ("2nd year - 25% tax savings on 40% depreciation", "₹", "46,872.00"),
            ("3rd year - 25% tax savings on 20% depreciation", "₹", "21,092.40"),
            ("**Total", "₹", "1,20,045.00"),
        ],
        [2.0, 0.2, 1.0]
    )
    right_cell.add_paragraph()
    
    create_table_section(
        right_cell, "Return On Investment (ROI) Calculation", 
        [
            ("Project Cost ex GST", "₹", "5,67,151.20"),
            ("Additional / Subsidy Benefits", "₹", "1,20,045.00"),
            ("Cost Via Grid", "₹", "2,31,376.32"),
            ("ROI in Years", "", "1.93"),
            ("Monthly Payments", "₹", "-"),
            ("**Total Plant cost inc Interest", "₹", "-"),
        ],
        [1.8, 0.4, 1.0]
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 25-Year Projection Table
    cols = ['Period', 'Unit\nGeneration', 'Cost via\nGrid', 'Solar Plant\nEMIs', 'O&M Cost', 'GSC\nCharges', 'AD Benefit', 'Savings via\nSolar']
    widths = [0.6, 0.8, 0.9, 0.7, 0.7, 0.7, 0.8, 0.9]
    proj_table = doc.add_table(rows=27, cols=8)
    proj_table.style = 'Table Grid'
    
    for i, w in enumerate(widths):
        for cell in proj_table.columns[i].cells:
            cell.width = Inches(w)
            
    # Header
    for i, text in enumerate(cols):
        format_cell(proj_table.rows[0].cells[i], text, bold=True, color_hex='FFFFFF', align='center', bg_color='002060', size=8)
        
    # Data Rows
    for r in range(1, 26):
        bg = 'DCE6F1' if (r-1) % 2 == 0 else 'FFFFFF'
        cells = proj_table.rows[r].cells
        format_cell(cells[0], f"Year {r}", align='center', bg_color=bg, size=7)
        format_cell(cells[1], "11,592.00", align='right', bg_color=bg, size=7)
        format_cell(cells[2], "₹ 2,08,656.00", align='right', bg_color=bg, size=7)
        format_cell(cells[3], "₹ -", align='right', bg_color=bg, size=7)
        format_cell(cells[4], "₹ 6,300.00", align='right', bg_color=bg, size=7)
        format_cell(cells[5], "₹ 22,720.32", align='right', bg_color=bg, size=7)
        format_cell(cells[6], "₹ 52,080.00", align='right', bg_color=bg, size=7)
        format_cell(cells[7], "₹ 2,83,456.32", align='right', bg_color=bg, size=7)
        
    # Totals Row
    cells = proj_table.rows[26].cells
    format_cell(cells[0], "Total over 25 years", bold=True, color_hex='FFFFFF', align='center', bg_color='002060', size=7)
    format_cell(cells[1], "2,69,863.33", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)
    format_cell(cells[2], "₹ 48,57,539.87", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)
    format_cell(cells[3], "₹ -", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)
    format_cell(cells[4], "₹ 2,16,886.76", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)
    format_cell(cells[5], "₹ 5,28,932.12", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)
    format_cell(cells[6], "₹ 1,20,044.40", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)
    format_cell(cells[7], "₹ 52,89,629.63", bold=True, color_hex='FFFFFF', align='right', bg_color='002060', size=7)
    
    doc.save(output_path)
    print(f"Saved to {output_path}")

if __name__ == '__main__':
    generate_evaluation_sheet('test_eval.docx')
