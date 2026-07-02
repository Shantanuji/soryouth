import docx

doc = docx.Document("scratch/template_patched.docx")
print("Number of paragraphs:", len(doc.paragraphs))
p = doc.paragraphs[99]
print("Patched Paragraph 99 Text:", p.text)
print("Patched Paragraph 99 runs count:", len(p.runs))
for r in p.runs:
    print(f"  Run: '{r.text}'")

# Search for any remaining Commercial Offer links
remaining = 0
for idx, p in enumerate(doc.paragraphs):
    p_xml = p._p.xml
    if 'LINK Excel' in p_xml and 'Commercial Offer' in p_xml:
        remaining += 1

print(f"Remaining Commercial Offer links: {remaining}")
