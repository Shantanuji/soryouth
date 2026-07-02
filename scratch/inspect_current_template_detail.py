import docx

doc = docx.Document("scratch/template_current.docx")
print("Number of paragraphs:", len(doc.paragraphs))
print("Number of tables:", len(doc.tables))
print("Number of sections:", len(doc.sections))

for idx, section in enumerate(doc.sections):
    print(f"\nSection {idx}:")
    header = section.header
    footer = section.footer
    print("  Header paragraphs:", len(header.paragraphs))
    print("  Header tables:", len(header.tables))
    print("  Footer paragraphs:", len(footer.paragraphs))
    print("  Footer tables:", len(footer.tables))

print("\n--- Inspecting First Table If Exists ---")
if doc.tables:
    table = doc.tables[0]
    print(f"Table 0 shape: {len(table.rows)} x {len(table.columns) if table.rows else 0}")
