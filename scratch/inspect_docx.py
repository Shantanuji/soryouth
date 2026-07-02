import docx
import sys

# Set standard output encoding to UTF-8
sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("test_eval.docx")
print("--- Paragraphs ---")
for i, p in enumerate(doc.paragraphs[:50]):
    if p.text.strip():
        print(f"{i}: {p.text}")

print("--- Tables ---")
for i, table in enumerate(doc.tables):
    print(f"\nTable {i}:")
    for r, row in enumerate(table.rows):
        row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        print(f"Row {r}: {row_text}")
