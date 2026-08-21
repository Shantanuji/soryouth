import docx

doc = docx.Document("api_test_output.docx")
print("Total native Word tables in generated docx:", len(doc.tables))
for idx, table in enumerate(doc.tables):
    print(f"\n--- Table #{idx+1} --- (Rows: {len(table.rows)}, Cols: {len(table.columns)})")
    for r in table.rows[:3]:
        cell_texts = [c.text.strip().replace('\n', ' ') for c in r.cells]
        print("  Row:", [s.encode('ascii', 'ignore').decode() for s in cell_texts])
