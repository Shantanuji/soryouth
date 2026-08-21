import os
import tempfile
import io
import base64
import subprocess
import platform
import shutil
import re
import uuid

try:
    import matplotlib
    matplotlib.use('Agg')  # Use non-interactive backend for servers
    from matplotlib.figure import Figure
    import matplotlib.ticker as mticker
except ImportError:
    matplotlib = None
    Figure = None
    mticker = None

from flask import Flask, request, jsonify
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import docx  # From python-docx
from html2image import Html2Image
import jinja2

class LenientUndefined(jinja2.Undefined):
    """
    Never fail on undefined/mismatched placeholders.
    Outputs the placeholder name as `{{ name }}` in the generated Word document
    so the user can visually locate and fix the placeholder without any generation error.
    """
    def _fail_with_undefined_error(self, *args, **kwargs):
        return f"{{{{ {self._undefined_name or 'undefined'} }}}}"
    
    def __str__(self):
        return f"{{{{ {self._undefined_name or 'undefined'} }}}}"
        
    def __iter__(self):
        return iter([])
        
    def __bool__(self):
        return False

if platform.system() == 'Windows':
    try:
        import pythoncom
        from docx2pdf import convert as convert_to_pdf
    except ImportError:
        pythoncom = None
        convert_to_pdf = None
else:
    pythoncom = None
    convert_to_pdf = None

app = Flask(__name__)


def safe_float(val, default=0.0):
    """Global helper to safely parse numerical strings into floats."""
    if val is None:
        return float(default)
    try:
        return float(str(val).replace(',', '').replace('₹', '').replace('%', '').strip())
    except (ValueError, TypeError):
        return float(default)


def format_indian_currency(val):
    """Formats numeric values to Indian style currency formatting (en-IN) with exactly 2 decimal places."""
    if val is None:
        return "0.00"
    val_str = str(val).strip()
    if not val_str or val_str.lower() == "none" or val_str.lower() == "n/a":
        return "0.00"
    
    # Strip any pre-existing formatting to get a clean raw numeric string
    cleaned = val_str.replace(',', '').replace('₹', '').replace('%', '').strip()
    if cleaned.startswith('-'):
        cleaned = cleaned[1:]
    if cleaned.startswith('+'):
        cleaned = cleaned[1:]
        
    try:
        float(cleaned)
    except ValueError:
        # Not a purely numeric value (e.g. "Included") - return as-is
        return val_str
        
    try:
        num = float(cleaned)
    except (ValueError, TypeError):
        return val_str
        
    neg = False
    if num < 0:
        neg = True
        num = abs(num)
        
    s = f"{num:.2f}"
    parts = s.split('.')
    int_part = parts[0]
    dec_part = parts[1]
    
    if len(int_part) <= 3:
        res = int_part
    else:
        last3 = int_part[-3:]
        remaining = int_part[:-3]
        groups = []
        while len(remaining) > 2:
            groups.append(remaining[-2:])
            remaining = remaining[:-2]
        if remaining:
            groups.append(remaining)
        groups.reverse()
        res = ",".join(groups) + "," + last3
        
    if neg:
        res = "-" + res
        
    return f"{res}.{dec_part}"


def get_html2image(size=(840, 1188), custom_flags=None):
    """Helper to instantiate Html2Image with explicitly resolved browser binary path on Linux."""
    if custom_flags is None:
        custom_flags = ['--no-sandbox', '--disable-gpu', '--force-device-scale-factor=3']
    
    browser_executable = None
    if platform.system() == 'Linux':
        for path in ['/usr/bin/google-chrome', '/usr/bin/google-chrome-stable', '/usr/bin/chromium-browser', '/usr/bin/chromium']:
            if os.path.exists(path):
                browser_executable = path
                break
    
    kwargs = {
        'size': size,
        'custom_flags': custom_flags
    }
    if browser_executable:
        kwargs['browser_executable'] = browser_executable
        
    return Html2Image(**kwargs)


def get_printable_width(doc):
    """
    Dynamically calculates 95% of the printable width of the DOCX template.
    Defaults to 6.5 inches if it cannot be determined.
    """
    try:
        section = doc.docx.sections[0]
        page_width = section.page_width.inches
        left_margin = section.left_margin.inches
        right_margin = section.right_margin.inches
        printable_width = page_width - left_margin - right_margin
        return Inches(printable_width * 0.95)
    except Exception:
        return Inches(6.5)


def create_monthly_generation_chart_b64(capacity_kw):
    """Generates monthly chart and returns base64 PNG or SVG HTML for rendering."""
    if not capacity_kw or capacity_kw <= 0:
        return None
        
    generation_per_day = capacity_kw * 4
    months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
    days_in_month = [31, 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]
    seasonal_factors = [0.95, 0.97, 1.10, 1.13, 1.14, 0.93, 0.75, 0.79, 0.87, 1.02, 1.00, 0.99]
    monthly_generation = [generation_per_day * days * factor for days, factor in zip(days_in_month, seasonal_factors)]

    if Figure is not None:
        try:
            fig = Figure(figsize=(9, 4.8), dpi=300)
            ax = fig.add_subplot(1, 1, 1)

            ax.bar(months, monthly_generation, color='#4F81BD', edgecolor='#385D8A', linewidth=1.5, width=0.6)
            ax.tick_params(axis='x', which='both', bottom=False, top=False, labelbottom=False)
            ax.tick_params(axis='y', colors='#595959', labelsize=12)
            
            ax.set_ylabel('Energy Produced (in kWh)', color='#595959', fontweight='bold', fontsize=14, labelpad=10)
            ax.set_xlim(-0.5, 11.5)
            
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['bottom'].set_visible(False)
            ax.spines['left'].set_visible(False)
            ax.grid(axis='y', linestyle='-', color='#D9D9D9', alpha=0.7)
            ax.set_axisbelow(True)
            
            cell_text = [[f"{int(val)}" for val in monthly_generation]]
            table = ax.table(cellText=cell_text, rowLabels=['Series1'], colLabels=months, loc='bottom', cellLoc='center', bbox=[0, -0.15, 1, 0.12])
            table.auto_set_font_size(False)
            table.set_fontsize(14)
            table.scale(1, 1.8)  # Normal padding
            
            for key, cell in table.get_celld().items():
                cell.set_edgecolor('#D9D9D9')
                if key[0] == 0:  # Header row
                    cell.set_text_props(weight='bold', color='#595959')
                if key[1] == -1: # Row label (Series1)
                    cell.set_text_props(weight='bold', color='#595959')

            fig.subplots_adjust(left=0.12, bottom=0.20, right=0.95, top=0.95)

            memfile = io.BytesIO()
            fig.savefig(memfile, format='png', transparent=True)
            memfile.seek(0)
            b64 = base64.b64encode(memfile.read()).decode('utf-8')
            return f"data:image/png;base64,{b64}"
        except Exception as e:
            print(f"Matplotlib monthly chart error: {e}")

    # Fallback to pure SVG / HTML
    max_val = max(monthly_generation) * 1.15 if monthly_generation else 1000
    svg_bars = []
    svg_table_cells = []
    for i, (m, val) in enumerate(zip(months, monthly_generation)):
        x = 55 + i * 66
        h = (val / max_val) * 180
        y = 220 - h
        svg_bars.append(f'<rect x="{x}" y="{y:.1f}" width="44" height="{h:.1f}" fill="#4F81BD" stroke="#385D8A" stroke-width="1.5" rx="3"/>')
        svg_bars.append(f'<text x="{x+22}" y="{y-6:.1f}" text-anchor="middle" font-size="11" fill="#333" font-weight="600">{int(val)}</text>')
        svg_table_cells.append(f'<td style="border:1px solid #D9D9D9; padding:4px 2px; text-align:center; font-size:11px; font-weight:600; color:#595959;">{m}<br/><span style="font-weight:normal;">{int(val)}</span></td>')

    return f'''
    <div style="font-family:sans-serif; width:880px; padding:10px; background:#fff;">
      <svg width="860" height="240" viewBox="0 0 860 240">
        <line x1="45" y1="220" x2="850" y2="220" stroke="#595959" stroke-width="1"/>
        <line x1="45" y1="20" x2="45" y2="220" stroke="#595959" stroke-width="1"/>
        <text x="20" y="120" transform="rotate(-90 20,120)" text-anchor="middle" fill="#595959" font-size="12" font-weight="bold">Energy Produced (kWh)</text>
        {''.join(svg_bars)}
      </svg>
      <table style="width:100%; border-collapse:collapse; margin-top:5px;">
        <tr>{''.join(svg_table_cells)}</tr>
      </table>
    </div>
    '''


def create_yearly_savings_chart_b64(capacity_kw, unit_rate):
    """Generates yearly chart and returns base64 PNG or SVG HTML for rendering."""
    if not capacity_kw or capacity_kw <= 0 or not unit_rate or unit_rate <= 0:
        return None

    initial_generation_per_year = capacity_kw * 4 * 365
    years = list(range(1, 31))
    savings = []
    
    current_generation = initial_generation_per_year
    current_rate = unit_rate

    for _ in years:
        savings.append(current_generation * current_rate)
        current_generation *= 0.996
        current_rate *= 1.02

    if Figure is not None:
        try:
            fig = Figure(figsize=(9, 4.8), dpi=300)
            ax = fig.add_subplot(1, 1, 1)

            ax.plot(years, savings, marker='o', linestyle='-', color='#F79646', linewidth=2.5, markersize=8)
            
            ax.set_ylabel('Projected Savings (in ₹)', color='#595959', fontweight='bold', fontsize=14, labelpad=10)
            ax.set_xlabel('Year', color='#595959', fontweight='bold', fontsize=14, labelpad=10)
            
            ax.set_xlim(0.8, 30.2) 
            
            ax.grid(True, which='both', linestyle='--', linewidth=0.5, color='#D9D9D9')
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['bottom'].set_color('#595959')
            ax.spines['left'].set_color('#595959')
            ax.tick_params(axis='both', colors='#595959', labelsize=12)
            if mticker is not None:
                ax.get_yaxis().set_major_formatter(mticker.FuncFormatter(lambda x, p: format(int(x), ',')))
            
            fig.subplots_adjust(left=0.15, bottom=0.15, right=0.95, top=0.95)

            memfile = io.BytesIO()
            fig.savefig(memfile, format='png', transparent=True)
            memfile.seek(0)
            b64 = base64.b64encode(memfile.read()).decode('utf-8')
            return f"data:image/png;base64,{b64}"
        except Exception as e:
            print(f"Matplotlib yearly chart error: {e}")

    # Fallback to pure SVG / HTML
    max_sav = max(savings) * 1.15 if savings else 1000000
    points = []
    svg_dots = []
    for i, (yr, sav) in enumerate(zip(years, savings)):
        x = 55 + i * 27
        y = 220 - (sav / max_sav) * 180
        points.append(f"{x:.1f},{y:.1f}")
        svg_dots.append(f'<circle cx="{x:.1f}" cy="{y:.1f}" r="3.5" fill="#F79646" stroke="#fff" stroke-width="1"/>')

    return f'''
    <div style="font-family:sans-serif; width:880px; padding:10px; background:#fff;">
      <svg width="860" height="240" viewBox="0 0 860 240">
        <line x1="45" y1="220" x2="850" y2="220" stroke="#595959" stroke-width="1"/>
        <line x1="45" y1="20" x2="45" y2="220" stroke="#595959" stroke-width="1"/>
        <text x="20" y="120" transform="rotate(-90 20,120)" text-anchor="middle" fill="#595959" font-size="12" font-weight="bold">Projected Savings (₹)</text>
        <polyline points="{' '.join(points)}" fill="none" stroke="#F79646" stroke-width="2.5"/>
        {''.join(svg_dots)}
        <text x="450" y="235" text-anchor="middle" fill="#595959" font-size="12" font-weight="bold">Years (1 - 30)</text>
      </svg>
    </div>
    '''


def create_combined_charts_page(doc, capacity_kw, unit_rate, target_width):
    """
    Generates an executive, print-quality solar performance page containing:
    1. Elegant Top Header: 'SOLAR PERFORMANCE & FINANCIAL ANALYSIS'
    2. Monthly Energy Generation Bar Chart (tall, integrated month pills at base, values on top, zero detached table)
    3. 30-Year Cumulative Financial Savings Chart (dynamic rising curve with milestone badges)
    Fills the entire A4 page majestically.
    """
    if not capacity_kw or capacity_kw <= 0:
        capacity_kw = 10
    if not unit_rate or unit_rate <= 0:
        unit_rate = 8.5

    generation_per_day = capacity_kw * 4
    months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
    days_in_month = [31, 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]
    seasonal_factors = [0.95, 0.97, 1.10, 1.13, 1.14, 0.93, 0.75, 0.79, 0.87, 1.02, 1.00, 0.99]
    monthly_gen = [int(generation_per_day * days * factor) for days, factor in zip(days_in_month, seasonal_factors)]
    annual_gen = sum(monthly_gen)
    
    max_m = max(monthly_gen) * 1.12 if monthly_gen else 1000
    
    # 1. Monthly SVG elements (Height: 480px)
    svg_bars = []
    baseline_y = 420
    chart_h = 350
    
    for i, (m, val) in enumerate(zip(months, monthly_gen)):
        x = 40 + i * 71.8
        h = (val / max_m) * chart_h
        y = baseline_y - h
        
        # Rounded gradient bar
        svg_bars.append(f'<rect x="{x}" y="{y:.1f}" width="54" height="{h:.1f}" fill="url(#barGrad)" rx="6"/>')
        # Value badge / text right above the bar
        svg_bars.append(f'<text x="{x+27}" y="{y-12:.1f}" text-anchor="middle" font-size="13.5" fill="#0F172A" font-weight="800">{val:,}</text>')
        # Integrated Month Pill attached right at the base of the bar
        svg_bars.append(f'<rect x="{x+2}" y="{baseline_y+10}" width="50" height="28" rx="14" fill="#EFF6FF" stroke="#BFDBFE" stroke-width="1.5"/>')
        svg_bars.append(f'<text x="{x+27}" y="{baseline_y+29}" text-anchor="middle" font-size="13" fill="#1D4ED8" font-weight="900">{m}</text>')

    # 2. 30-Year Cumulative Savings (Height: 480px)
    initial_gen = capacity_kw * 4 * 365
    years = list(range(1, 31))
    savings = []
    cg = initial_gen
    cr = unit_rate
    for _ in years:
        savings.append(cg * cr)
        cg *= 0.996
        cr *= 1.02
        
    total_savings_30 = sum(savings)
    min_s = min(savings) * 0.70
    max_s = max(savings) * 1.10
    span_s = max_s - min_s if max_s > min_s else 1
    
    points = []
    svg_dots = []
    baseline_y_sav = 420
    chart_h_sav = 350
    
    for i, (yr, sav) in enumerate(zip(years, savings)):
        x = 55 + i * 27.2
        ratio = (sav - min_s) / span_s
        y = baseline_y_sav - ratio * chart_h_sav
        points.append(f"{x:.1f},{y:.1f}")
        
        if yr in [1, 5, 10, 15, 20, 25, 30]:
            val_str = f"&#8377; {sav/100000:.1f} L" if sav < 10000000 else f"&#8377; {sav/10000000:.2f} Cr"
            svg_dots.append(f'<circle cx="{x:.1f}" cy="{y:.1f}" r="7" fill="#EA580C" stroke="#FFFFFF" stroke-width="3"/>')
            # Pill badge for savings values
            svg_dots.append(f'<rect x="{x-36:.1f}" y="{y-34:.1f}" width="72" height="24" rx="12" fill="#FFF7ED" stroke="#FED7AA" stroke-width="1.5"/>')
            svg_dots.append(f'<text x="{x:.1f}" y="{y-17:.1f}" text-anchor="middle" font-size="12.5" fill="#C2410C" font-weight="900">{val_str}</text>')
            # Year pill at bottom
            svg_dots.append(f'<rect x="{x-26:.1f}" y="{baseline_y_sav+10}" width="52" height="28" rx="14" fill="#F8FAFC" stroke="#CBD5E1" stroke-width="1.5"/>')
            svg_dots.append(f'<text x="{x:.1f}" y="{baseline_y_sav+29}" text-anchor="middle" font-size="12" fill="#334155" font-weight="900">Yr {yr}</text>')
        else:
            svg_dots.append(f'<circle cx="{x:.1f}" cy="{y:.1f}" r="3.5" fill="#F97316" opacity="0.8"/>')

    area_points = [f"55,{baseline_y_sav}"] + points + [f"{55 + 29 * 27.2:.1f},{baseline_y_sav}"]

    css = """
    <style>
      * { box-sizing: border-box; margin: 0; padding: 0; }
      html, body {
        overflow: hidden !important;
        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
        background: #FFFFFF;
        width: 880px !important;
        height: 1140px !important;
        padding: 8px 12px !important;
        display: flex;
        flex-direction: column;
        justify-content: space-between;
        color: #0F172A;
      }
      ::-webkit-scrollbar {
        display: none !important;
        width: 0px !important;
        height: 0px !important;
      }
      .main-title-box {
        background: linear-gradient(135deg, #0F3B66 0%, #1B4D75 100%);
        border-radius: 10px;
        padding: 10px 18px;
        display: flex;
        justify-content: space-between;
        align-items: center;
        color: #FFFFFF;
        box-shadow: 0 3px 8px rgba(15, 59, 102, 0.12);
      }
      .main-title {
        font-size: 18px;
        font-weight: 900;
        letter-spacing: 0.4px;
        text-transform: uppercase;
      }
      .main-subtitle {
        font-size: 11px;
        opacity: 0.9;
        font-weight: 500;
        margin-top: 2px;
      }
      .title-metric {
        text-align: right;
        background: rgba(255, 255, 255, 0.15);
        padding: 4px 12px;
        border-radius: 6px;
        border: 1px solid rgba(255, 255, 255, 0.25);
      }
      .title-metric-val {
        font-size: 15px;
        font-weight: 900;
      }
      .title-metric-lbl {
        font-size: 9.5px;
        opacity: 0.85;
        text-transform: uppercase;
        letter-spacing: 0.4px;
      }
      .chart-card {
        background: #FFFFFF;
        border: 1.2px solid #E2E8F0;
        border-radius: 10px;
        padding: 10px 16px;
        display: flex;
        flex-direction: column;
        box-shadow: 0 3px 8px rgba(0, 0, 0, 0.03);
        flex: 1;
        margin-top: 8px;
        justify-content: space-between;
        overflow: hidden;
      }
      .card-header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        border-bottom: 1.2px solid #F1F5F9;
        padding-bottom: 6px;
      }
      .card-title {
        font-size: 15px;
        font-weight: 800;
        color: #0F2942;
        letter-spacing: 0.2px;
      }
      .card-badges {
        display: flex;
        gap: 8px;
      }
      .badge {
        background: #EFF6FF;
        color: #1D4ED8;
        font-size: 11.5px;
        font-weight: 800;
        padding: 3px 10px;
        border-radius: 16px;
        border: 1.2px solid #BFDBFE;
      }
      .badge-orange {
        background: #FFF7ED;
        color: #C2410C;
        border: 1.2px solid #FED7AA;
      }
    </style>
    """

    html = f'''
    <!DOCTYPE html>
    <html>
    <head>
    <meta charset="utf-8"/>
    {css}
    </head>
    <body>

      <!-- Top Page Header -->
      <div class="main-title-box">
        <div>
          <div class="main-title">Solar Performance &amp; Financial Analysis</div>
          <div class="main-subtitle">Detailed 12-Month Generation Profile &amp; 30-Year Cumulative Financial Growth</div>
        </div>
        <div class="title-metric">
          <div class="title-metric-val">{capacity_kw} kWp</div>
          <div class="title-metric-lbl">System Capacity</div>
        </div>
      </div>

      <!-- Chart 1: Monthly Energy Generation -->
      <div class="chart-card">
        <div class="card-header">
          <div class="card-title">Monthly Energy Generation (Year 1)</div>
          <div class="card-badges">
            <div class="badge">Annual Total: {annual_gen:,} kWh</div>
            <div class="badge" style="background:#F0FDF4; color:#166534; border-color:#BBF7D0;">Daily Avg: {int(annual_gen/365):,} Units</div>
          </div>
        </div>
        
        <svg width="850" height="420" viewBox="0 0 900 480">
          <defs>
            <linearGradient id="barGrad" x1="0" y1="0" x2="0" y2="1">
              <stop offset="0%" stop-color="#3B82F6"/>
              <stop offset="50%" stop-color="#2563EB"/>
              <stop offset="100%" stop-color="#1D4ED8"/>
            </linearGradient>
          </defs>
          <!-- Grid lines -->
          <line x1="35" y1="90" x2="890" y2="90" stroke="#F1F5F9" stroke-width="1.5"/>
          <line x1="35" y1="190" x2="890" y2="190" stroke="#F1F5F9" stroke-width="1.5"/>
          <line x1="35" y1="290" x2="890" y2="290" stroke="#F1F5F9" stroke-width="1.5"/>
          <line x1="35" y1="{baseline_y}" x2="890" y2="{baseline_y}" stroke="#94A3B8" stroke-width="1.5"/>
          
          <text x="14" y="220" transform="rotate(-90 14,220)" text-anchor="middle" fill="#64748B" font-size="12" font-weight="800">Energy (kWh)</text>
          {''.join(svg_bars)}
        </svg>
      </div>

      <!-- Chart 2: 30-Year Cumulative Savings -->
      <div class="chart-card">
        <div class="card-header">
          <div class="card-title">30-Year Projected Cumulative Financial Savings</div>
          <div class="card-badges">
            <div class="badge badge-orange">Total 30-Yr Savings: &#8377; {total_savings_30/10000000:.2f} Cr</div>
            <div class="badge badge-orange" style="background:#FEF2F2; color:#991B1B; border-color:#FECACA;">Grid Tariff: &#8377; {unit_rate}/kWh</div>
          </div>
        </div>
        
        <svg width="850" height="420" viewBox="0 0 900 480">
          <defs>
            <linearGradient id="areaGrad" x1="0" y1="0" x2="0" y2="1">
              <stop offset="0%" stop-color="#F97316" stop-opacity="0.40"/>
              <stop offset="60%" stop-color="#F97316" stop-opacity="0.12"/>
              <stop offset="100%" stop-color="#F97316" stop-opacity="0.01"/>
            </linearGradient>
          </defs>
          <!-- Grid lines -->
          <line x1="35" y1="90" x2="890" y2="90" stroke="#F1F5F9" stroke-width="1.5"/>
          <line x1="35" y1="190" x2="890" y2="190" stroke="#F1F5F9" stroke-width="1.5"/>
          <line x1="35" y1="290" x2="890" y2="290" stroke="#F1F5F9" stroke-width="1.5"/>
          <line x1="35" y1="{baseline_y_sav}" x2="890" y2="{baseline_y_sav}" stroke="#94A3B8" stroke-width="1.5"/>
          
          <text x="14" y="220" transform="rotate(-90 14,220)" text-anchor="middle" fill="#64748B" font-size="12" font-weight="800">Cumulative Savings (&#8377;)</text>
          
          <polygon points="{' '.join(area_points)}" fill="url(#areaGrad)"/>
          <polyline points="{' '.join(points)}" fill="none" stroke="#EA580C" stroke-width="3.8" stroke-linecap="round" stroke-linejoin="round"/>
          {''.join(svg_dots)}
        </svg>
      </div>

    </body>
    </html>
    '''
    
    hti = get_html2image(size=(880, 1140), custom_flags=['--no-sandbox', '--disable-gpu', '--hide-scrollbars', '--force-device-scale-factor=2.5'])
    
    out_dir = tempfile.gettempdir()
    hti.output_path = out_dir
    hti.temp_path = out_dir

    filename = f'combined_charts_{uuid.uuid4().hex}.png'
    temp_path = os.path.join(hti.output_path, filename)
    
    try:
        hti.screenshot(html_str=html, save_as=filename)
        if os.path.exists(temp_path):
            # Scale to fit standard A4 printable bounds perfectly without spilling to next page in LibreOffice/Word
            return InlineImage(doc, temp_path, width=Inches(6.4))
        else:
            return None
    except Exception as e:
        print(f"Error generating combined charts screenshot: {e}")
        return None


# BACKWARD COMPATIBILITY: Keep old functions intact but unused
def create_monthly_generation_chart(doc, capacity_kw, target_width):
    if not capacity_kw or capacity_kw <= 0:
        return None
    generation_per_day = capacity_kw * 4
    months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
    days_in_month = [31, 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]
    seasonal_factors = [0.95, 0.97, 1.10, 1.13, 1.14, 0.93, 0.75, 0.79, 0.87, 1.02, 1.00, 0.99]
    monthly_generation = [generation_per_day * days * factor for days, factor in zip(days_in_month, seasonal_factors)]

    fig = Figure(figsize=(8, 4.5), dpi=300)
    ax = fig.add_subplot(1, 1, 1)
    ax.bar(months, monthly_generation, color='#4F81BD', edgecolor='#385D8A', linewidth=1.5, width=0.6)
    ax.tick_params(axis='x', which='both', bottom=False, top=False, labelbottom=False)
    ax.tick_params(axis='y', colors='#595959', labelsize=10)
    ax.set_ylabel('Energy Produced (in kWh)', color='#595959', fontweight='bold', fontsize=12, labelpad=10)
    ax.set_title('Monthly Projection', color='#595959', fontweight='bold', fontsize=16, pad=15)
    ax.set_xlim(-0.5, 11.5)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.grid(axis='y', linestyle='-', color='#D9D9D9', alpha=0.7)
    ax.set_axisbelow(True)
    
    cell_text = [[f"{int(val)}" for val in monthly_generation]]
    table = ax.table(cellText=cell_text, rowLabels=['Series1'], colLabels=months, loc='bottom', cellLoc='center', bbox=[0, -0.22, 1, 0.12])
    table.auto_set_font_size(False)
    table.set_fontsize(13)
    for key, cell in table.get_celld().items():
        cell.set_edgecolor('#D9D9D9')

    fig.subplots_adjust(left=0.12, bottom=0.30, right=0.95, top=0.90)

    memfile = io.BytesIO()
    fig.savefig(memfile, format='png')
    memfile.seek(0)
    return InlineImage(doc, memfile, width=target_width)


def create_yearly_savings_chart(doc, capacity_kw, unit_rate, target_width):
    if not capacity_kw or capacity_kw <= 0 or not unit_rate or unit_rate <= 0:
        return None
    initial_generation_per_year = capacity_kw * 4 * 365
    years = list(range(1, 31))
    savings = []
    current_generation = initial_generation_per_year
    current_rate = unit_rate
    for _ in years:
        savings.append(current_generation * current_rate)
        current_generation *= 0.996
        current_rate *= 1.02

    fig = Figure(figsize=(8, 4.5), dpi=300)
    ax = fig.add_subplot(1, 1, 1)
    ax.plot(years, savings, marker='o', linestyle='-', color='#F79646', linewidth=2.5, markersize=6)
    ax.set_ylabel('Projected Savings (in ₹)', color='#595959', fontweight='bold', fontsize=12)
    ax.set_xlabel('Year', color='#595959', fontweight='bold', fontsize=12)
    ax.set_title('Projected Yearly Savings for 30 Years', color='#595959', fontweight='bold', fontsize=16, pad=15)
    ax.set_xlim(0.8, 30.2) 
    ax.grid(True, which='both', linestyle='--', linewidth=0.5, color='#D9D9D9')
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_color('#595959')
    ax.spines['left'].set_color('#595959')
    ax.tick_params(axis='both', colors='#595959', labelsize=10)
    ax.get_yaxis().set_major_formatter(mticker.FuncFormatter(lambda x, p: format(int(x), ',')))
    
    fig.subplots_adjust(left=0.12, bottom=0.30, right=0.95, top=0.90)

    memfile = io.BytesIO()
    fig.savefig(memfile, format='png')
    memfile.seek(0)
def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_borders(cell, top="0080C0", bottom="0080C0", left="0080C0", right="0080C0", sz="12", val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders_elm = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{top}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{left}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{bottom}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{right}"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    margins_elm = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(margins_elm)

def add_native_balance_of_system_table(container, context):
    """
    Generates a native, fully editable Word table inside any Document or Subdoc
    with exact blue styling, borders, shading, font weights, and dynamic values.
    """
    import math

    capacity = safe_float(context.get('capacity', context.get('project_size', 0)))
    mod_watt = safe_float(context.get('module_wattage'), 600)
    if mod_watt <= 0:
        mod_watt = 600
        
    calc_mod_qty = int(math.ceil((capacity * 1000) / mod_watt)) if capacity > 0 else 0
    passed_mod_qty = safe_float(context.get('module_qty', context.get('moduleQty', 0)))
    module_qty = str(int(passed_mod_qty)) if passed_mod_qty > 0 else (str(calc_mod_qty) if calc_mod_qty > 0 else "14")

    mod_type = str(context.get('module_type', 'Topcon Bifacial')).strip()
    dcr_stat = str(context.get('dcr_status', 'DCR')).strip()
    mod_spec_fallback = f"Rayzon Solar {mod_type} {dcr_stat} {int(mod_watt)} Wp"
    module_spec = str(context.get('module_spec', context.get('module_details', context.get('module_description', mod_spec_fallback))))

    inv_kw = safe_float(context.get('inverter_rating', context.get('inverter_kw', capacity)))
    inv_kw_fmt = f"{int(inv_kw) if inv_kw == int(inv_kw) else inv_kw:.1f} kW" if inv_kw > 0 else "8 kW"
    inv_spec_fallback = f"Growatt/Sungrow {inv_kw_fmt}"
    inverter_spec = str(context.get('inverter_spec', context.get('inverter_details', context.get('inverter_description', inv_spec_fallback))))

    inv_qty_val = int(safe_float(context.get('inverter_qty', context.get('inverterQty', 1))))
    inv_qty_val = max(1, inv_qty_val)
    inverter_qty_nos = f"{inv_qty_val} Nos"

    acdb_qty_val = int(safe_float(context.get('acdb_qty', context.get('acdb_dcdb_qty', inv_qty_val))))
    dcdb_qty_val = int(safe_float(context.get('dcdb_qty', context.get('acdb_dcdb_qty', inv_qty_val))))
    earthing_qty_val = int(safe_float(context.get('earthing_kit_qty', context.get('earthing_qty', inv_qty_val * 3))))
    la_qty_val = int(safe_float(context.get('la_kit_qty', context.get('la_qty', inv_qty_val))))

    dcdb_qty_nos = f"{dcdb_qty_val} No" if dcdb_qty_val == 1 else f"{dcdb_qty_val} Nos"
    acdb_qty_nos = f"{acdb_qty_val} No" if acdb_qty_val == 1 else f"{acdb_qty_val} Nos"
    earthing_kit_qty_nos = f"{earthing_qty_val} Nos"
    la_kit_qty_nos = f"{la_qty_val} No" if la_qty_val == 1 else f"{la_qty_val} Nos"

    p = container.add_paragraph()
    run = p.add_run("Balance of System")
    run.font.name = 'Segoe UI'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 59, 96) # #0B3B60
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(14)

    rows_data = [
        ("Solar Modules", "PV Solar Module", module_qty, module_spec),
        ("Inverter", inverter_spec, inverter_qty_nos, inverter_spec),
        ("Mounting Structure", "RCC Rooftop", "1 Lot", "GI Sqaure Pipe for Columns & Rafter/Strut Channel for Panel Mounting 15-18ft"),
        ("DC Cable", "4 sq.mm", "As per design", "Polycab 1.1 kV Standard EN 50618 / IEC 62930 UV Resistant - Yes"),
        ("AC Cable", "As per Design", "As per design", "XLPE Insulated Armed Polycab Voltage Rating 1.1 kV Standard - IS 7098"),
        ("DCDB", "As per design", dcdb_qty_nos, "DC Isolator, DC SPD Type-II, String Fuses Enclosure - IP65 Weatherproof"),
        ("ACDB", "As per design", acdb_qty_nos, "MCCB/MCB, AC SPD Type-II Enclosure - IP65 Weatherproof"),
        ("Earthing Kit", "Chemical Earthing", earthing_kit_qty_nos, "Electrode Size - 17.2 mm Copper Bonded / GI Electrode Earth Resistance less than 5 Ohms Compliance IS 3043"),
        ("Lightning Arrester", "Standard", la_kit_qty_nos, "IS/IEC Standards"),
        ("Connectors", "MC4 compatible", "As required", "1500 V DC Protection Class - IP68"),
        ("Monitoring", "RMS App-based", "1 Set", "Real-Time Generation, Fault Alerts, Historical Data Analysis"),
        ("Tags", "Aluminum Engraved", "1 Set", "Aluminum Engraved Identification Tags"),
        ("Fire Extinguishers", "As per compliance", "1 Set", "As per MNRE / Electrical Compliance Requirements"),
    ]

    table = container.add_table(rows=len(rows_data) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_widths = [Inches(1.3), Inches(1.5), Inches(0.9), Inches(2.8)]

    # Header row
    hdr_cells = table.rows[0].cells
    headers = ["Component", "Details", "Qty", "Specifications"]
    for i, h_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.width = col_widths[i]
        set_cell_background(cell, "1B4D75")
        set_cell_borders(cell, top="0080C0", bottom="0080C0", left="0080C0", right="0080C0", sz="12")
        set_cell_margins(cell, top=160, bottom=160, left=180, right=180)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        hp = cell.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hrun = hp.add_run(h_text)
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(10.5)
        hrun.font.bold = True
        hrun.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r_idx, row_tuple in enumerate(rows_data, start=1):
        row_cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row_tuple):
            cell = row_cells[c_idx]
            cell.width = col_widths[c_idx]
            set_cell_background(cell, "EAF4FC")
            set_cell_borders(cell, top="0080C0", bottom="0080C0", left="0080C0", right="0080C0", sz="10")
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            cp = cell.paragraphs[0]
            if c_idx == 2:
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            crun = cp.add_run(str(val))
            crun.font.name = 'Calibri'
            crun.font.size = Pt(9.5)
            crun.font.color.rgb = RGBColor(11, 59, 96)
            if c_idx == 0:
                crun.font.bold = True

    return table


def generate_balance_of_system_png(context):
    """
    Generates an ultra-high-resolution PNG screenshot of the Balance of System page
    with exact Soryouth branding, component list, quantities, and specifications
    derived directly from the proposal/database context.
    """
    import math

    capacity = safe_float(context.get('capacity', context.get('project_size', 0)))
    mod_watt = safe_float(context.get('module_wattage'), 600)
    if mod_watt <= 0:
        mod_watt = 600
        
    calc_mod_qty = int(math.ceil((capacity * 1000) / mod_watt)) if capacity > 0 else 0
    passed_mod_qty = safe_float(context.get('module_qty', context.get('moduleQty', 0)))
    module_qty = str(int(passed_mod_qty)) if passed_mod_qty > 0 else (str(calc_mod_qty) if calc_mod_qty > 0 else "14")

    mod_type = str(context.get('module_type', 'Topcon Bifacial')).strip()
    dcr_stat = str(context.get('dcr_status', 'DCR')).strip()
    mod_spec_fallback = f"Rayzon Solar {mod_type} {dcr_stat} {int(mod_watt)} Wp"
    module_spec = str(context.get('module_spec', context.get('module_details', context.get('module_description', mod_spec_fallback))))

    inv_kw = safe_float(context.get('inverter_rating', context.get('inverter_kw', capacity)))
    inv_kw_fmt = f"{int(inv_kw) if inv_kw == int(inv_kw) else inv_kw:.1f} kW" if inv_kw > 0 else "8 kW"
    inv_spec_fallback = f"Growatt/Sungrow {inv_kw_fmt}"
    inverter_spec = str(context.get('inverter_spec', context.get('inverter_details', context.get('inverter_description', inv_spec_fallback))))

    inv_qty_val = int(safe_float(context.get('inverter_qty', context.get('inverterQty', 1))))
    inv_qty_val = max(1, inv_qty_val)
    inverter_qty_nos = f"{inv_qty_val} Nos"

    acdb_qty_val = int(safe_float(context.get('acdb_qty', context.get('acdb_dcdb_qty', inv_qty_val))))
    dcdb_qty_val = int(safe_float(context.get('dcdb_qty', context.get('acdb_dcdb_qty', inv_qty_val))))
    earthing_qty_val = int(safe_float(context.get('earthing_kit_qty', context.get('earthing_qty', inv_qty_val * 3))))
    la_qty_val = int(safe_float(context.get('la_kit_qty', context.get('la_qty', inv_qty_val))))

    dcdb_qty_nos = f"{dcdb_qty_val} No" if dcdb_qty_val == 1 else f"{dcdb_qty_val} Nos"
    acdb_qty_nos = f"{acdb_qty_val} No" if acdb_qty_val == 1 else f"{acdb_qty_val} Nos"
    earthing_kit_qty_nos = f"{earthing_qty_val} Nos"
    la_kit_qty_nos = f"{la_qty_val} No" if la_qty_val == 1 else f"{la_qty_val} Nos"

    out_dir = tempfile.gettempdir()
    logo_filename = f"soryouth_logo_{uuid.uuid4().hex[:8]}.png"
    logo_dest_path = os.path.join(out_dir, logo_filename)

    script_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.abspath(os.path.join(script_dir, '..', '..'))

    possible_logo_paths = [
        os.path.join(project_root, 'public', 'assets', 'images', 'logo-icon.png'),
        os.path.join(project_root, 'public', 'assets', 'images', 'soryouth-logo.png'),
        os.path.join(project_root, 'public', 'assets', 'images', 'logo.png'),
        os.path.join(project_root, 'public', 'assets', 'images', 'logo-dark.png'),
        os.path.join(project_root, 'public', 'logo.png'),
        os.path.abspath('public/assets/images/logo-icon.png'),
        os.path.abspath('public/assets/images/soryouth-logo.png'),
        os.path.abspath('public/assets/images/logo.png'),
    ]
    found_logo = None
    for p in possible_logo_paths:
        if os.path.exists(p):
            found_logo = p
            break

    logo_html = ""
    if found_logo:
        try:
            shutil.copyfile(found_logo, logo_dest_path)
            logo_html = f'<img src="{logo_filename}" style="height: 76px; max-width: 220px; object-fit: contain; margin-right: 25px;" />'
        except Exception as ex:
            print(f"Error copying logo for Chromium rendering: {ex}")
            logo_html = ""

    bos_css = """
    <style>
        * {
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }
        ::-webkit-scrollbar {
            display: none;
        }
        body {
            font-family: 'Segoe UI', Calibri, Arial, sans-serif;
            background: #ffffff;
            width: 840px;
            height: 960px;
            padding: 12px 20px;
            margin: 0;
            overflow: hidden;
            color: #0B3B60;
            display: flex;
            flex-direction: column;
            justify-content: space-between;
        }
        .header-container {
            display: flex;
            align-items: center;
            margin-bottom: 10px;
            padding-bottom: 0px;
        }
        .title {
            font-family: 'Georgia', 'Cambria', 'Times New Roman', serif;
            font-size: 32px;
            font-weight: 700;
            color: #0F3B66;
            letter-spacing: 0;
            line-height: 1;
        }
        table {
            width: 100%;
            flex-grow: 1;
            border-collapse: collapse;
            border: 2px solid #0080C0;
            font-size: 12px;
        }
        th {
            background-color: #1B4D75;
            color: #ffffff;
            font-weight: 700;
            font-size: 13px;
            padding: 8px 10px;
            text-align: center;
            border: 1.5px solid #0080C0;
        }
        td {
            padding: 6px 10px;
            border: 1.5px solid #0080C0;
            background-color: #EAF4FC;
            color: #0B3B60;
            vertical-align: middle;
            font-size: 12px;
            line-height: 1.25;
        }
        .component-name {
            font-weight: 700;
            color: #0B3B60;
        }
        .center-col {
            text-align: center;
            white-space: nowrap;
        }
    </style>
    """

    html = f'''
    <html>
    <head>
    {bos_css}
    </head>
    <body>
        <div class="header-container">
            {logo_html}
            <div class="title">Balance of System</div>
        </div>

        <table>
            <thead>
                <tr>
                    <th style="width: 18%;">Component</th>
                    <th style="width: 22%;">Details</th>
                    <th style="width: 15%;">Qty</th>
                    <th style="width: 45%;">Specifications</th>
                </tr>
            </thead>
            <tbody>
                <tr>
                    <td class="component-name">Solar Modules</td>
                    <td>PV Solar Module</td>
                    <td class="center-col">{module_qty}</td>
                    <td>{module_spec}</td>
                </tr>
                <tr>
                    <td class="component-name">Inverter</td>
                    <td>{inverter_spec}</td>
                    <td class="center-col">{inverter_qty_nos}</td>
                    <td>{inverter_spec}</td>
                </tr>
                <tr>
                    <td class="component-name">Mounting Structure</td>
                    <td>RCC Rooftop</td>
                    <td class="center-col">1 Lot</td>
                    <td>GI Sqaure Pipe for Columns & Rafter/Strut Channel for Panel Mounting 15-18ft</td>
                </tr>
                <tr>
                    <td class="component-name">DC Cable</td>
                    <td>4 sq.mm</td>
                    <td class="center-col">As per design</td>
                    <td>Polycab 1.1 kV Standard EN 50618 / IEC 62930 UV Resistant - Yes</td>
                </tr>
                <tr>
                    <td class="component-name">AC Cable</td>
                    <td>As per Design</td>
                    <td class="center-col">As per design</td>
                    <td>XLPE Insulated Armed Polycab Voltage Rating 1.1 kV Standard - IS 7098</td>
                </tr>
                <tr>
                    <td class="component-name">DCDB</td>
                    <td>As per design</td>
                    <td class="center-col">{dcdb_qty_nos}</td>
                    <td>DC Isolator, DC SPD Type-II, String Fuses Enclosure - IP65 Weatherproof</td>
                </tr>
                <tr>
                    <td class="component-name">ACDB</td>
                    <td>As per design</td>
                    <td class="center-col">{acdb_qty_nos}</td>
                    <td>MCCB/MCB, AC SPD Type-II Enclosure - IP65 Weatherproof</td>
                </tr>
                <tr>
                    <td class="component-name">Earthing Kit</td>
                    <td>Chemical Earthing</td>
                    <td class="center-col">{earthing_kit_qty_nos}</td>
                    <td>Electrode Size - 17.2 mm Copper Bonded / GI Electrode Earth Resistance less than 5 Ohms Compliance IS 3043</td>
                </tr>
                <tr>
                    <td class="component-name">Lightning Arrester</td>
                    <td>Standard</td>
                    <td class="center-col">{la_kit_qty_nos}</td>
                    <td>IS/IEC Standards</td>
                </tr>
                <tr>
                    <td class="component-name">Connectors</td>
                    <td>MC4 compatible</td>
                    <td class="center-col">As required</td>
                    <td>1500 V DC Protection Class - IP68</td>
                </tr>
                <tr>
                    <td class="component-name">Monitoring</td>
                    <td>RMS App-based</td>
                    <td class="center-col">1 Set</td>
                    <td>Real-Time Generation, Fault Alerts, Historical Data Analysis</td>
                </tr>
                <tr>
                    <td class="component-name">Tags</td>
                    <td>Aluminum Engraved</td>
                    <td class="center-col">1 Set</td>
                    <td>Aluminum Engraved Identification Tags</td>
                </tr>
                <tr>
                    <td class="component-name">Fire Extinguishers</td>
                    <td>As per compliance</td>
                    <td class="center-col">1 Set</td>
                    <td>As per MNRE / Electrical Compliance Requirements</td>
                </tr>
            </tbody>
        </table>
    </body>
    </html>
    '''

    hti = get_html2image(size=(840, 960), custom_flags=['--no-sandbox', '--disable-gpu', '--force-device-scale-factor=3'])
    out_dir = tempfile.gettempdir()
    hti.output_path = out_dir
    hti.temp_path = out_dir

    filename = f'balance_of_system_{uuid.uuid4().hex}.png'
    temp_path = os.path.join(out_dir, filename)
    
    try:
        res = hti.screenshot(html_str=html, save_as=filename)
        target_path = None
        if res and isinstance(res, list) and len(res) > 0 and os.path.exists(res[0]):
            target_path = res[0]
        else:
            for p in [temp_path, os.path.join(os.getcwd(), filename), filename]:
                if os.path.exists(p) and os.path.getsize(p) > 0:
                    target_path = p
                    break

        if target_path and os.path.exists(target_path):
            with open(target_path, 'rb') as f:
                return target_path, f.read()
    except Exception as e:
        print(f"Error generating Balance of System PNG: {e}")
    return None, None


def style_cell_custom(cell, text, bg="FFFFFF", fg=(0,0,0), bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=6.8, border_color="0080C0"):
    set_cell_background(cell, bg)
    set_cell_borders(cell, top=border_color, bottom=border_color, left=border_color, right=border_color, sz="4")
    set_cell_margins(cell, top=10, bottom=10, left=30, right=30)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    run.font.name = 'Segoe UI'
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*fg)
    return run


def add_native_project_scope_table(container, context=None):
    """
    Generates a native, fully editable Microsoft Word table for Project's Scope.
    Matches exact design of Screenshot 1 and fits cleanly on 1 A4 page.
    """
    scope_data = [
        ("1", "Supply of Items", "PV Solar components as per BOM & BOQ", "Yes", "-"),
        ("", "Supply of Items", "Roof / Superstructure / Elevated Structure", "Yes", "-"),
        ("", "Supply of Items", "Data Logger", "Yes", "-"),
        ("2", "Freight", "Transportation", "Yes", "-"),
        ("", "Freight", "Packing & Handling", "Yes", "-"),
        ("3", "Material Handling at Site", "Safe unloading of material from vehicle", "Yes", "-"),
        ("", "Material Handling at Site", "Lifting & shifting of materials at site", "Yes", "-"),
        ("", "Material Handling at Site", "Parking space / permission for vehicles", "-", "Yes"),
        ("4", "Material Storage", "Safe & secure storage space at site", "-", "Yes"),
        ("5", "Roof", "Waterproofing", "-", "Yes"),
        ("", "Roof", "NOC / permission for roof usage", "-", "Yes"),
        ("", "Roof", "Safe access to installation site", "-", "Yes"),
        ("", "Roof", "Clearance of installation area", "-", "Yes"),
        ("6", "Liaisoning", "Net metering liaisoning with utility", "Yes", "-"),
        ("7", "Installation & Commissioning (I&C)", "Design & Engineering of Solar PV System", "Yes", "-"),
        ("", "Installation & Commissioning (I&C)", "Tools, tackles & installation consumables", "Yes", "-"),
        ("", "Installation & Commissioning (I&C)", "Solar PV system installation services", "Yes", "-"),
        ("", "Installation & Commissioning (I&C)", "Lightning Arrestor (LA)", "Yes", "-"),
        ("", "Installation & Commissioning (I&C)", "Earthing & system safety", "Yes", "-"),
        ("", "Installation & Commissioning (I&C)", "Civil works for installation", "Yes", "-"),
        ("", "Installation & Commissioning (I&C)", "Testing & Commissioning", "Yes", "-"),
    ]

    p = container.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Project's Scope")
    run.font.name = 'Georgia'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 59, 102) # #0F3B66
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)

    table = container.add_table(rows=len(scope_data) + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_widths = [Inches(0.6), Inches(1.8), Inches(3.2), Inches(0.8), Inches(0.8)]
    headers = ["Sr. No.", "Category", "Scope Item", "Soryouth's", "Client's"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1B4D75")
        set_cell_borders(hdr_cells[i], top="0080C0", bottom="0080C0", left="0080C0", right="0080C0", sz="4")
        set_cell_margins(hdr_cells[i], top=15, bottom=15, left=30, right=30)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hp = hdr_cells[i].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(0)
        hrun = hp.add_run(h)
        hrun.font.name = 'Segoe UI'
        hrun.font.size = Pt(7.5)
        hrun.font.bold = True
        hrun.font.color.rgb = RGBColor(255, 255, 255)

    for r_idx, (sr, cat, item, sor, cli) in enumerate(scope_data, start=1):
        bg = "EAF4FC" if r_idx % 2 == 1 else "FFFFFF"
        r_cells = table.rows[r_idx].cells
        trPr = r_cells[0]._tc.getparent().get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

        for c_idx, w in enumerate(col_widths):
            r_cells[c_idx].width = w

        style_cell_custom(r_cells[0], sr, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=6.8, bold=bool(sr))
        style_cell_custom(r_cells[1], cat, bg=bg, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=6.8, bold=True, fg=(11,59,96))
        style_cell_custom(r_cells[2], item, bg=bg, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=6.8, fg=(11,59,96))
        style_cell_custom(r_cells[3], sor, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=6.8, bold=(sor=="Yes"))
        style_cell_custom(r_cells[4], cli, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=6.8, bold=(cli=="Yes"))

    return table


def add_native_terms_and_conditions_table(container, context=None):
    """
    Generates a native, fully editable Microsoft Word table for Term's & Conditions.
    Matches exact design of Screenshot 2 and fits cleanly on 1 A4 page.
    """
    validity_days = 15
    if context:
        try:
            val_in = context.get('validityDays', context.get('validity_days', 15))
            validity_days = int(val_in) if val_in else 15
        except Exception:
            validity_days = 15

    validity_text = f"Proposal valid for {validity_days} days. Price subject to change if scope, specifications, or statutory requirements change. Excludes permits, panel upgrades, roof repairs, trenching unless specified."

    terms_data = [
        ("1", "Pricing & Validity", validity_text),
        ("2", "Taxes", "GST shall be charged extra at actuals as per applicable rates."),
        ("3", "Freight & Logistics", "Inclusive of transportation, packing, forwarding, and unloading at site."),
        ("4", "Technical Site Assessment", "Site verification post-LOI/PO. Any variation may impact cost. No warranty on roof or electrical panel condition."),
        ("5", "Delivery & Timeline", "Delivery within 8 weeks from advance payment and PO confirmation. Subject to extension due to delays, approvals, force majeure, or client-side dependencies."),
        ("6", "Approvals & Documentation", "Seller will assist as agent for approvals. Timelines depend on authorities; delays not attributable to Seller."),
        ("7", "Proposal & IP", "All documents remain Seller's IP. Cannot be reused without consent. Financial estimates are indicative only."),
        ("8", "Insurance", "Seller responsible till commissioning. Post-delivery and post-commissioning insurance responsibility lies with Purchaser."),
        ("9", "Warranty (General)", "Equipment and services warranted against defects as per defined terms."),
        ("10", "Warranty (Equipment)", "Modules: 12 yrs product + 30 yrs performance. Inverter: 5 yrs. BOS: 2 yrs"),
        ("11", "Warranty Exclusions", "Excludes wear & tear, misuse, unauthorized changes, environmental damage, theft, force majeure."),
        ("12", "Pre-Warranty Damages", "Damages due to negligence or external factors before warranty start shall be billed extra."),
    ]

    p = container.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Term's & Conditions")
    run.font.name = 'Georgia'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 59, 102) # #0F3B66
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)

    table = container.add_table(rows=len(terms_data) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_widths = [Inches(0.6), Inches(2.1), Inches(4.5)]
    headers = ["Sr. No.", "Section", "Details"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1B4D75")
        set_cell_borders(hdr_cells[i], top="0080C0", bottom="0080C0", left="0080C0", right="0080C0", sz="4")
        set_cell_margins(hdr_cells[i], top=15, bottom=15, left=30, right=30)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hp = hdr_cells[i].paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(0)
        hrun = hp.add_run(h)
        hrun.font.name = 'Segoe UI'
        hrun.font.size = Pt(7.5)
        hrun.font.bold = True
        hrun.font.color.rgb = RGBColor(255, 255, 255)

    for r_idx, (sr, sec, details) in enumerate(terms_data, start=1):
        bg = "EAF4FC" if r_idx % 2 == 1 else "FFFFFF"
        r_cells = table.rows[r_idx].cells
        trPr = r_cells[0]._tc.getparent().get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

        for c_idx, w in enumerate(col_widths):
            r_cells[c_idx].width = w

        style_cell_custom(r_cells[0], sr, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=6.5, bold=True)
        style_cell_custom(r_cells[1], sec, bg=bg, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=6.5, bold=True, fg=(11,59,96))
        style_cell_custom(r_cells[2], details, bg=bg, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=6.5, fg=(11,59,96))

    return table


def create_native_project_scope_subdoc(doc, context=None):
    try:
        subdoc = doc.new_subdoc()
        add_native_project_scope_table(subdoc, context)
        return subdoc
    except Exception as e:
        print(f"Error creating native project scope subdoc: {e}")
        return None


def create_native_terms_conditions_subdoc(doc, context=None):
    try:
        subdoc = doc.new_subdoc()
        add_native_terms_and_conditions_table(subdoc, context)
        return subdoc
    except Exception as e:
        print(f"Error creating native terms and conditions subdoc: {e}")
        return None


def replace_scope_and_terms_static_images(doc, raw_context):
    """
    Scans the Word document for static drawings ONLY for Project Scope and Term's & Conditions pages,
    removes the static drawing XML tags AND their containing paragraph, inserting Native Editable Microsoft Word Tables in-place.
    Cleans up empty trailing paragraphs to guarantee zero blank orphan pages.
    All other pages remain 100% untouched.
    """
    image_rel_map = {}
    for rel_id, rel in doc.part.rels.items():
        ref_lower = str(rel.target_ref).lower()
        if 'image4.png' in ref_lower or 'image4.jpeg' in ref_lower:
            image_rel_map['scope'] = rel_id
        elif 'image8.png' in ref_lower or 'image8.jpeg' in ref_lower or 'image9.png' in ref_lower or 'image9.jpeg' in ref_lower:
            image_rel_map['terms'] = rel_id

    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }

    temp_builder_doc = docx.Document()

    for p in list(doc.paragraphs):
        p_xml = p._element
        drawings = p_xml.findall('.//w:drawing', ns)
        for dwg in drawings:
            blips = dwg.findall('.//a:blip', ns)
            for b in blips:
                embed_id = b.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_id and embed_id == image_rel_map.get('scope'):
                    print(f"Replacing static drawing {embed_id} (Project's Scope image4.png) with native Word Table in-place...")
                    tbl = add_native_project_scope_table(temp_builder_doc, raw_context)
                    p_parent = p_xml.getparent()
                    if p_parent is not None:
                        idx = p_parent.index(p_xml)
                        p_parent.insert(idx, tbl._element)
                        p_parent.remove(p_xml)

                elif embed_id and embed_id == image_rel_map.get('terms'):
                    print(f"Replacing static drawing {embed_id} (Term's & Conditions image) with native Word Table in-place...")
                    tbl = add_native_terms_and_conditions_table(temp_builder_doc, raw_context)
                    p_parent = p_xml.getparent()
                    if p_parent is not None:
                        idx = p_parent.index(p_xml)
                        p_parent.insert(idx, tbl._element)
                        p_parent.remove(p_xml)

    # Clean up empty trailing paragraphs after inserted tables to prevent blank orphan pages
    for table in doc.tables:
        tbl_xml = table._element
        curr = tbl_xml.getnext()
        while curr is not None and curr.tag.endswith('p'):
            p_elem = curr
            text = p_elem.text.strip() if p_elem.text else ''
            p_drawings = p_elem.findall('.//w:drawing', ns)
            p_blips = p_elem.findall('.//a:blip', ns)
            next_sibling = p_elem.getnext()
            
            if text == '' and len(p_drawings) == 0 and len(p_blips) == 0:
                parent = p_elem.getparent()
                parent.remove(p_elem)
                curr = next_sibling
            else:
                break


def create_native_balance_of_system_subdoc(doc, context):
    try:
        subdoc = doc.new_subdoc()
        add_native_balance_of_system_table(subdoc, context)
        return subdoc
    except Exception:
        return None


def create_balance_of_system_page(doc, context, target_width):
    path, blob = generate_balance_of_system_png(context)
    if path and os.path.exists(path):
        return InlineImage(doc, path, width=target_width)
    return None


def create_capex_evaluation_sheet(doc, context, target_width):
    """
    Generates a print-quality HTML Evaluation Sheet and securely renders it into 
    an ultra-high-resolution Native A4 image using Chromium.
    """
    name       = str(context.get('name', 'N/A'))
    location   = str(context.get('location', 'N/A'))
    capacity   = safe_float(context.get('capacity'))
    rate_pw    = safe_float(context.get('rate_per_watt'))
    cost_pkw   = rate_pw * 1000
    base_amt   = safe_float(context.get('base_amount'), cost_pkw * capacity)
    final_amt  = safe_float(context.get('final_amount'), safe_float(context.get('total_project_cost_inc_gst'), base_amt * 1.09))
    gst_amt    = max(0.0, final_amt - base_amt)
    gst_pct    = (gst_amt / base_amt * 100) if base_amt > 0 else 9.0
    unit_rate  = safe_float(context.get('unit_rate'), safe_float(context.get('grid_tariff_per_unit')))
    subsidy    = safe_float(context.get('subsidy_amount'))
    gen_yr     = safe_float(context.get('generation_per_year'), capacity * 4 * 345)
    savings_yr = safe_float(context.get('savings_per_year'), gen_yr * unit_rate)

    client_type = str(context.get('client_type', context.get('clientType', 'Other'))).strip().lower()
    unit_rate_val = safe_float(context.get('unit_rate'), safe_float(context.get('grid_tariff_per_unit')))
    is_business = ('commercial' in client_type) or ('industrial' in client_type) or ('industr' in client_type) or ('industry' in client_type) or ('factory' in client_type) or ('business' in client_type) or ('corporate' in client_type) or (unit_rate_val <= 12 and unit_rate_val > 0 and 'housing' not in client_type and 'bungalow' not in client_type and 'individual' not in client_type)

    if is_business:
        subsidy = 0.0
        ad1_val = safe_float(context.get('ad_benefit_year1'), 0.0)
        ad1 = ad1_val if ad1_val > 0 else (base_amt * 0.40 * 0.25)
        
        ad2_val = safe_float(context.get('ad_benefit_year2'), 0.0)
        ad2 = ad2_val if ad2_val > 0 else ((base_amt - ad1) * 0.40 * 0.25)
        
        ad3_val = safe_float(context.get('ad_benefit_year3'), 0.0)
        ad3 = ad3_val if ad3_val > 0 else ((base_amt - ad1 - ad2) * 0.20 * 0.25)
        
        total_ad_val = safe_float(context.get('total_ad_benefit'), 0.0)
        total_ad = total_ad_val if total_ad_val > 0 else (ad1 + ad2 + ad3)
    else:
        ad1 = 0.0
        ad2 = 0.0
        ad3 = 0.0
        total_ad = 0.0

    om_pkw   = safe_float(context.get('om_cost_per_kw'), 750)
    om_base  = safe_float(context.get('total_om_cost'), capacity * om_pkw)
    
    net_inv = final_amt - subsidy - total_ad
    roi_years = safe_float(context.get('roi_in_years'), net_inv / savings_yr if savings_yr > 0 else 0)

    ad_benefits_html = f'''
                <!-- AD Benefits -->
                <table>
                    <tr><td colspan="3" class="table-title">Accelerated Depreciation Benefits</td></tr>
                    <tr class="blue-row"><td class="label-col">1st year - 25% tax savings on 40% depreciation</td><td class="label-col">₹</td><td class="val-col">{ad1:,.2f}</td></tr>
                    <tr class="white-row"><td class="label-col">2nd year - 25% tax savings on 40% depreciation</td><td class="label-col">₹</td><td class="val-col">{ad2:,.2f}</td></tr>
                    <tr class="blue-row"><td class="label-col">3rd year - 25% tax savings on 20% depreciation</td><td class="label-col">₹</td><td class="val-col">{ad3:,.2f}</td></tr>
                    <tr class="white-row" style="font-weight:bold"><td class="label-col">Total</td><td class="label-col">₹</td><td class="val-col">{total_ad:,.2f}</td></tr>
                </table>
    '''
    ad_roi_row_html = f'''
                    <tr class="blue-row"><td class="label-col">Accelerated Depreciation Benefits</td><td class="label-col">₹</td><td class="val-col">{total_ad:,.2f}</td></tr>
    '''
        
    subsidy_roi_row_html = f'''
                    <tr class="white-row" style="font-weight:bold"><td class="label-col">subsidy</td><td class="label-col">₹</td><td class="val-col">{subsidy:,.2f}</td></tr>
    '''

    html = f'''
    <html>
    <head>
    <style>
        * {{
            box-sizing: border-box;
        }}
        ::-webkit-scrollbar {{
            display: none;
        }}
        body {{
            font-family: 'Segoe UI', Calibri, sans-serif;
            background: white;
            padding: 10px 20px;
            width: 840px;
            height: 1188px;
            margin: 0;
            overflow: hidden;
            color: #000;
        }}
        h2 {{
            text-align: center;
            color: #002060;
            font-style: italic;
            font-weight: bold;
            font-size: 24px;
            margin-top: 10px;
            margin-bottom: 15px;
            letter-spacing: 0.5px;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin-bottom: 15px;
            font-size: 11px;
            border: 1px solid #222;
        }}
        th, td {{
            border: 1px solid #222;
            padding: 3px 5px;
        }}
        th {{
            background-color: #002060;
            color: white;
            text-align: center;
            font-weight: bold;
        }}
        .table-title {{
            background-color: #002060 !important;
            color: white !important;
            text-align: center !important;
            font-weight: bold !important;
            padding: 4px;
            font-size: 11.5px;
            letter-spacing: 0.3px;
        }}
        .green-row td {{
            background-color: #92D050;
            font-weight: bold;
        }}
        .blue-row td {{
            background-color: #DCE6F1;
        }}
        .white-row td {{
            background-color: #FFFFFF;
        }}
        .col-layout {{
            display: flex;
            justify-content: space-between;
            gap: 15px;
        }}
        .left-col, .right-col {{
            width: 50%;
        }}
        .val-col {{ text-align: right; }}
        .label-col {{ text-align: left; font-weight: 500; }}
        
        .arrows-section {{
            text-align: center;
            margin: 5px 0 10px 0;
            display: flex;
            justify-content: center;
            align-items: center;
            gap: 20px;
            color: #002060;
            font-weight: bold;
            font-size: 14px;
        }}
        .arrow {{
            font-size: 26px;
            line-height: 1;
        }}
        
        .proj-table th {{
            font-size: 9.5px;
            padding: 3px;
        }}
        .proj-table td {{
            font-size: 9.5px;
            padding: 2.5px 4px;
        }}
        .notes {{
            font-size: 9.5px;
            color: #333;
            margin-top: 10px;
            line-height: 1.3;
        }}
        ul {{
            margin: 3px 0 0 15px;
            padding: 0;
        }}
    </style>
    </head>
    <body>
        <h2>Evaluation Sheet for Capex Model</h2>
        
        <div class="col-layout">
            <div class="left-col">
                <!-- Project Specification -->
                <table>
                    <tr><td colspan="3" class="table-title">Project Specification</td></tr>
                    <tr class="blue-row"><td class="label-col">Client Name</td><td class="val-col" colspan="2">{name}</td></tr>
                    <tr class="white-row"><td class="label-col">Project Location</td><td class="val-col" colspan="2">{location}</td></tr>
                    <tr class="blue-row"><td class="label-col">Project Size (KW)</td><td class="val-col" colspan="2">{capacity:,.2f}</td></tr>
                    <tr class="green-row"><td class="label-col">Cost per KW ex GST</td><td class="label-col">₹</td><td class="val-col">{cost_pkw:,.2f}</td></tr>
                    <tr class="blue-row"><td class="label-col">Project Cost ex GST</td><td class="label-col">₹</td><td class="val-col">{base_amt:,.2f}</td></tr>
                    <tr class="white-row"><td class="label-col">GST @ {gst_pct:.1f}%</td><td class="label-col">₹</td><td class="val-col">{gst_amt:,.2f}</td></tr>
                    <tr class="green-row"><td class="label-col">Total Project Cost inc GST</td><td class="label-col">₹</td><td class="val-col">{final_amt:,.2f}</td></tr>
                    <tr class="white-row"><td class="label-col">Client's Current Grid Tariff /Unit</td><td class="label-col">₹</td><td class="val-col">{unit_rate:,.2f}</td></tr>
                </table>
                
                <!-- Plant Performance -->
                <table>
                    <tr><td colspan="3" class="table-title">Plant Performance</td></tr>
                    <tr class="blue-row"><td class="label-col">Average Annual Generation (KW)</td><td class="val-col" colspan="2">{gen_yr:,.2f}</td></tr>
                    <tr class="white-row"><td class="label-col">Average Monthly Generation (KW)</td><td class="val-col" colspan="2">{gen_yr / 12:,.2f}</td></tr>
                    <tr class="blue-row"><td class="label-col">Estimated Year-on-Year degradation</td><td class="val-col" colspan="2">0.70 - 0.80%</td></tr>
                </table>
                
                <!-- O&M Cost -->
                <table>
                    <tr><td colspan="3" class="table-title">O&M Cost</td></tr>
                    <tr class="blue-row"><td class="label-col">Cost per KW of maintenance</td><td class="label-col">₹</td><td class="val-col">{om_pkw:,.2f}</td></tr>
                    <tr class="white-row"><td class="label-col">Total Cost per year</td><td class="label-col">₹</td><td class="val-col">{om_base:,.2f}</td></tr>
                    <tr class="blue-row"><td class="label-col">Yearly escalation in O&M Cost</td><td class="val-col" colspan="2">3.00%</td></tr>
                </table>
            </div>
            
            <div class="right-col">
                {ad_benefits_html}
                
                <!-- ROI Calculation -->
                <table>
                    <tr><td colspan="3" class="table-title">Return On Investment (ROI) Calculation</td></tr>
                    <tr class="white-row"><td class="label-col">Project Cost ex GST</td><td class="label-col">₹</td><td class="val-col">{base_amt:,.2f}</td></tr>
                    {ad_roi_row_html}
                    <tr class="white-row"><td class="label-col">Cost Via Grid</td><td class="label-col">₹</td><td class="val-col">{savings_yr * 25:,.2f}</td></tr>
                    <tr class="blue-row"><td class="label-col">ROI in Years</td><td class="val-col" colspan="2" style="text-align:right;">{roi_years:,.2f}</td></tr>
                    <tr class="white-row"><td class="label-col">Monthly Payments</td><td class="label-col">₹</td><td class="val-col">-</td></tr>
                    <tr class="blue-row"><td class="label-col">Total Plant cost inc Interest</td><td class="label-col">₹</td><td class="val-col">-</td></tr>
                    {subsidy_roi_row_html}
                </table>
            </div>
        </div>
        
        <div class="arrows-section">
            <div class="arrow">⬇</div>
            <div>Savings Projection</div>
            <div class="arrow">⬇</div>
        </div>
        
        <table class="proj-table">
            <tr>
                <th>Period</th>
                <th>Unit<br/>Generation</th>
                <th>Cost via<br/>Grid</th>
                <th>Solar Plant<br/>EMIs</th>
                <th>O&M Cost</th>
                <th>GSC<br/>Charges</th>
                <th>AD Benefit</th>
                <th>Savings via<br/>Solar</th>
            </tr>
    '''
    
    evaluation_sheet = context.get('evaluationSheet')
    yr_data = []
    sum_gen, sum_grid, sum_om, sum_gsc, sum_ad, sum_sav = 0, 0, 0, 0, 0, 0
    
    if evaluation_sheet and isinstance(evaluation_sheet, list) and len(evaluation_sheet) > 0:
        for row in evaluation_sheet:
            y = int(row.get('year', 0))
            gen = float(row.get('generation', 0))
            grid = float(row.get('gridCost', 0))
            om = float(row.get('omCost', 0))
            gsc = float(row.get('gscCharges', round(gen * 1.96, 2)))
            row_ad = float(row.get('adBenefit', 0))
            if is_business:
                ad = row_ad if row_ad > 0 else ([ad1, ad2, ad3][y - 1] if y <= 3 else 0.0)
            else:
                ad = 0.0
            sav = (grid + ad + gsc - om)
            yr_data.append((y, gen, grid, om, gsc, ad, sav))
    else:
        for y in range(1, 26):
            gen  = gen_yr * (0.992 ** (y - 1))
            grid = gen * unit_rate
            om   = 0.0 if y == 1 else om_base * (1.03 ** (y - 2))
            gsc  = gen * 1.96
            ad_val = ([ad1, ad2, ad3][y - 1] if y <= 3 else 0.0) if is_business else 0.0
            sav  = (grid + ad_val + gsc - om)
            yr_data.append((y, gen, grid, om, gsc, ad_val, sav))
            
    for r in yr_data:
        y, gen, grid, om, gsc, ad, sav = r
        sum_gen += gen
        sum_grid += grid
        sum_om += om
        sum_gsc += gsc
        sum_ad += ad
        sum_sav += sav
        
        bg_class = "blue-row" if (y - 1) % 2 == 0 else "white-row"
        
        om_str = f"₹ {om:,.2f}" if om > 0 else "₹ -"
        gsc_str = f"₹ {gsc:,.2f}" if gsc > 0 else "₹ -"
        ad_str = f"₹ {ad:,.2f}" if ad > 0 else "₹ -"
        
        html += f'''
            <tr class="{bg_class}">
                <td style="text-align:center">Year {y}</td>
                <td class="val-col">{gen:,.2f}</td>
                <td class="val-col">₹ {grid:,.2f}</td>
                <td class="val-col">₹ -</td>
                <td class="val-col">{om_str}</td>
                <td class="val-col">{gsc_str}</td>
                <td class="val-col">{ad_str}</td>
                <td class="val-col" style="font-weight:500;">₹ {sav:,.2f}</td>
            </tr>
        '''
        
    html += f'''
            <tr style="background-color: #002060; color: white;">
                <td class="table-title">Total over {len(yr_data)} years</td>
                <td class="table-title val-col">{sum_gen:,.3f}</td>
                <td class="table-title val-col">₹ {sum_grid:,.2f}</td>
                <td class="table-title val-col">₹ -</td>
                <td class="table-title val-col">₹ {sum_om:,.2f}</td>
                <td class="table-title val-col">₹ {sum_gsc:,.2f}</td>
                <td class="table-title val-col">₹ {sum_ad:,.2f}</td>
                <td class="table-title val-col">₹ {sum_sav:,.2f}</td>
            </tr>
        </table>
        
        <div class="notes">
            <u>The estimates provided in the above table are subject to the following conditions:</u>
            <ul>
                <li>The generation numbers are estimated as per current weather data and will change as per actual weather conditions.</li>
                <li>A standard generation reduction of 0.60% year-on-year is applied for further accuracy.</li>
                <li>The savings shown are based on continuous consumption of the generated power by the client.</li>
                <li>Any change in government net-metering policy or otherwise may affect the savings, and are out of our control.</li>
                <li>Any O&M cost incurred during the plant life will be extra, which isn't calculated in the table.</li>
                <li>Furthermore, any unexpected damage due to natural disaster is not factored in the calculations.</li>
            </ul>
        </div>
    </body>
    </html>
    '''
    
    # Native 3x rendering utilizing device scale factor for perfect 300 DPI A4 mapping
    hti = get_html2image(size=(840, 1188), custom_flags=['--no-sandbox', '--disable-gpu', '--force-device-scale-factor=3'])
    
    out_dir = tempfile.gettempdir()
    hti.output_path = out_dir
    hti.temp_path = out_dir

        
    filename = f'capex_{uuid.uuid4().hex}.png'
    temp_path = os.path.join(hti.output_path, filename)
    
    try:
        hti.screenshot(html_str=html, save_as=filename)
        if os.path.exists(temp_path):
            return InlineImage(doc, temp_path, width=target_width)
        else:
            print(f"Warning: html2image failed to create screenshot at {temp_path}.")
            return None
    except Exception as e:
        print(f"Error generating capex sheet screenshot: {e}")
        return None


@app.route('/extract-placeholders', methods=['POST'])
def extract_placeholders():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    if file:
        try:
            document = docx.Document(io.BytesIO(file.read()))
            placeholders = set()
            pattern = re.compile(r'\{\{([^}]+)\}\}')
            
            for para in document.paragraphs:
                matches = pattern.findall(para.text)
                for match in matches:
                    placeholders.add(f"{{{{{match}}}}}")
            
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            matches = pattern.findall(para.text)
                            for match in matches:
                                placeholders.add(f"{{{{{match}}}}}")

            return jsonify({"success": True, "placeholders": sorted(list(placeholders))})
        except Exception as e:
            return jsonify({"error": f"Error processing file: {str(e)}"}), 500
    return jsonify({"error": "File processing failed"}), 500


@app.route('/generate', methods=['POST'])
def generate_proposal():
    original_mplconfigdir = os.environ.get('MPLCONFIGDIR')
    temp_dir = tempfile.mkdtemp()
    
    try:
        os.environ['MPLCONFIGDIR'] = temp_dir
        payload = request.get_json()
        if not payload:
            return jsonify({"error": "Invalid JSON payload"}), 400
        
        template_full_path = payload.get('template_path')
        context = payload.get('data')
        
        if not template_full_path or not context:
            return jsonify({"error": "Missing 'template_path' or 'data' in payload"}), 400
        try:
            doc = DocxTemplate(template_full_path)
            doc.init_docx()
        except Exception as tpl_err:
            print(f"Warning: Failed to load template from {template_full_path}: {tpl_err}. Falling back to default master template.")
            master_candidates = [
                'public/uploads/templates/kapex_fixed_data.docx',
                'dist/kapex-fixed-data.docx',
                'public/uploads/templates/kapex-fixed-data-official.docx'
            ]
            doc = None
            for mc in master_candidates:
                if os.path.exists(mc):
                    try:
                        doc = DocxTemplate(mc)
                        doc.init_docx()
                        print(f"Successfully loaded fallback master template from {mc}")
                        break
                    except Exception:
                        pass
            if doc is None:
                raise tpl_err

        raw_context = dict(context) if isinstance(context, dict) else {}
        sub_amt = safe_float(context.get('subsidy_amount'))
        add_sub = safe_float(context.get('additional_subsidy'), safe_float(context.get('additional_subsidy_benefits')))
        total_sub = safe_float(context.get('total_subsidy_amount'), sub_amt + add_sub)
        if total_sub == 0.0 and (sub_amt > 0.0 or add_sub > 0.0):
            total_sub = sub_amt + add_sub

        total_ad = safe_float(context.get('total_ad_benefit'), safe_float(context.get('totalAdBenefit')))
        display_subsidy_or_ad = total_sub if total_sub > 0.0 else total_ad

        raw_context.update({
            'name':               str(context.get('name', 'N/A')),
            'location':           str(context.get('location', 'N/A')),
            'capacity':           safe_float(context.get('capacity')),
            'rate_per_watt':      safe_float(context.get('rate_per_watt')),
            'base_amount':        safe_float(context.get('base_amount')),
            'final_amount':       safe_float(context.get('final_amount')),
            'cgst_amount':        safe_float(context.get('cgst_amount')),
            'sgst_amount':        safe_float(context.get('sgst_amount')),
            'subsidy_amount':     display_subsidy_or_ad,
            'central_subsidy_amount': sub_amt if sub_amt > 0.0 else total_sub,
            'additional_subsidy_benefits': add_sub,
            'additional_subsidy': add_sub,
            'total_subsidy_amount': total_sub,
            'unit_rate':          safe_float(context.get('unit_rate')),
            'grid_tariff_per_unit': safe_float(context.get('unit_rate')),
            'generation_per_year':  safe_float(context.get('generation_per_year')),
            'savings_per_year':     safe_float(context.get('savings_per_year')),
            'evaluationSheet':    context.get('evaluationSheet'),
            'ad_benefit_year1':   context.get('ad_benefit_year1'),
            'ad_benefit_year2':   context.get('ad_benefit_year2'),
            'ad_benefit_year3':   context.get('ad_benefit_year3'),
            'total_ad_benefit':   context.get('total_ad_benefit'),
            'om_cost_per_kw':     context.get('om_cost_per_kw'),
            'total_om_cost':      context.get('total_om_cost'),
            'roi_in_years':       context.get('roi_in_years'),
        })

        capacity_kw = raw_context['capacity']
        unit_rate_val = raw_context['unit_rate']

        try:
            undeclared = doc.get_undeclared_template_variables()
        except Exception:
            undeclared = set()
            
        # Dynamically evaluate the maximum printable width for responsive rendering
        max_printable_width = get_printable_width(doc)

        # Ensure render context has all raw_context values
        context.update(raw_context)

        # -------------------------------------------------------------
        # ALWAYS render combined page and evaluation sheet so tag run-splitting never drops them
        if capacity_kw and capacity_kw > 0:
            combined_page_image = create_combined_charts_page(doc, capacity_kw, unit_rate_val, max_printable_width)
            if combined_page_image:
                context['combined_charts_page'] = combined_page_image
                context['monthly_generation_chart'] = combined_page_image
                context['yearly_savings_chart'] = combined_page_image

        capex_sheet_image = create_capex_evaluation_sheet(doc, raw_context, max_printable_width)
        if capex_sheet_image:
            context['capex_evaluation_sheet'] = capex_sheet_image
        # -------------------------------------------------------------
        
        bos_placeholders = ['balance_of_system', 'balance_of_system_page', 'balance_of_system_table', 'bos_table', 'balance_system', 'bos']
        bos_rendered = False
        
        bos_png_path, bos_png_bytes = generate_balance_of_system_png(raw_context)

        # 1. If template contains explicit BOS placeholders (e.g. {{balance_of_system}}), inject InlineImage or Subdoc
        if bos_png_path and os.path.exists(bos_png_path):
            bos_inline = InlineImage(doc, bos_png_path, width=max_printable_width)
            for key in bos_placeholders:
                context[key] = bos_inline
            bos_rendered = True
        else:
            bos_subdoc = create_native_balance_of_system_subdoc(doc, raw_context)
            if bos_subdoc:
                for key in bos_placeholders:
                    context[key] = bos_subdoc
                bos_rendered = True

        # Consistently format all financial values to en-IN style currency format (exactly 2 decimal places) for template rendering
        currency_keys = [
            'base_amount', 'final_amount', 'cgst_amount', 'sgst_amount', 'subsidy_amount',
            'central_subsidy_amount', 'additional_subsidy_benefits', 'additional_subsidy',
            'total_subsidy_amount', 'net_amount_after_subsidy', 'net_investment', 'netInvestment',
            'subtotal', 'project_cost_ex_gst', 'gst_amount', 'total_project_cost_inc_gst',
            'savings_per_year', 'total_om_cost', 'ad_benefit_year1', 'ad_benefit_year2',
            'ad_benefit_year3', 'total_ad_benefit', 'project_cost_ex_gst_roi', 'cost_via_grid',
            'cost_per_kw', 'rate_per_watt'
        ]
        for key in currency_keys:
            if key in context:
                context[key] = format_indian_currency(context[key])

        # Safeguard: Use LenientUndefined so any unknown, mismatched, or missing placeholder
        # in the template never causes a fatal error. Missing placeholders will display cleanly
        # as `{{ placeholder_name }}` or default value in the rendered output, allowing the user
        # to identify and fix issues in their template without breaking proposal generation.
        try:
            template_vars = doc.get_undeclared_template_variables()
            for v in template_vars:
                if v not in context:
                    # Provide an empty string or let LenientUndefined preserve {{ v }}
                    pass
        except Exception as e:
            print(f"Warning analyzing undeclared variables: {e}")

        jinja_env = jinja2.Environment(undefined=LenientUndefined, autoescape=False)
        try:
            doc.render(context, jinja_env=jinja_env)
        except Exception as render_ex:
            print(f"Primary render warning: {render_ex}. Retrying with empty fallback context...")
            try:
                # Fallback: fill all undeclared variables with empty strings and retry
                try:
                    for v in doc.get_undeclared_template_variables():
                        if v not in context:
                            context[v] = ""
                except Exception:
                    pass
                doc.render(context)
            except Exception as final_ex:
                print(f"Final render error: {final_ex}")
                # Save whatever was processed

            
        temp_docx_path = os.path.join(temp_dir, 'output.docx')
        
        # Clean trailing empty paragraphs at the end of the rendered document so Word file has zero trailing blank pages
        try:
            while len(doc.paragraphs) > 0:
                last_p = doc.paragraphs[-1]
                drawings = last_p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                if (not last_p.text or last_p.text.strip() == '') and len(drawings) == 0:
                    parent = last_p._element.getparent()
                    parent.remove(last_p._element)
                else:
                    break
        except Exception:
            pass

        doc.save(temp_docx_path)

        temp_pdf_path = os.path.join(temp_dir, 'output.pdf')
        
        try:
            # Generate Image-PDF: Render DOCX pages -> PIL Images -> Merge into crisp 150 DPI PDF
            raw_pdf_path = os.path.join(temp_dir, 'raw.pdf')
            if platform.system() == 'Windows':
                pythoncom.CoInitialize()
                try:
                    convert_to_pdf(temp_docx_path, raw_pdf_path)
                finally:
                    pythoncom.CoUninitialize()
            else:
                cmd = ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', temp_dir, temp_docx_path]
                subprocess.run(cmd, check=True)
                converted_name = os.path.splitext(os.path.basename(temp_docx_path))[0] + '.pdf'
                generated_pdf = os.path.join(temp_dir, converted_name)
                if os.path.exists(generated_pdf) and generated_pdf != raw_pdf_path:
                    os.replace(generated_pdf, raw_pdf_path)

            # Convert raw PDF pages to high-res PIL Images and compile into final PDF
            import pypdfium2 as pdfium
            from PIL import ImageStat
            
            raw_pdf = pdfium.PdfDocument(raw_pdf_path)
            page_images = []
            for i, page in enumerate(raw_pdf):
                pil_img = page.render(scale=2.0).to_pil().convert('RGB')
                stat = ImageStat.Stat(pil_img)
                is_blank = all(m > 254.95 for m in stat.mean) and all(s < 0.05 for s in stat.stddev)
                if is_blank:
                    print(f"[Image-PDF] Pruned 100% white blank page {i+1}")
                    continue
                page_images.append(pil_img)

            if page_images:
                page_images[0].save(temp_pdf_path, save_all=True, append_images=page_images[1:], resolution=150.0)
                print(f"[Image-PDF] Successfully compiled {len(page_images)}-page Image-PDF at {temp_pdf_path}")
            else:
                os.replace(raw_pdf_path, temp_pdf_path)

        except Exception as e:
            print(f"[Image-PDF] Image-PDF pipeline fallback error: {e}")
            if os.path.exists(os.path.join(temp_dir, 'raw.pdf')):
                os.replace(os.path.join(temp_dir, 'raw.pdf'), temp_pdf_path)
            elif not os.path.exists(temp_pdf_path):
                return jsonify({"error": f"PDF conversion failed. Error: {e}"}), 500

        # PDF Post-Processor: Verify generated PDF integrity
        try:
            import pypdf
            reader = pypdf.PdfReader(temp_pdf_path)
            print(f"[PDF Processor] Successfully verified PDF generation. Total pages: {len(reader.pages)}")
        except Exception as p_err:
            print(f"[PDF Processor] Notice: {p_err}")

        with open(temp_docx_path, 'rb') as f_docx, open(temp_pdf_path, 'rb') as f_pdf:
            docx_buffer = f_docx.read()
            pdf_buffer = f_pdf.read()

        docx_b64 = base64.b64encode(docx_buffer).decode('utf-8')
        pdf_b64  = base64.b64encode(pdf_buffer).decode('utf-8')
        
        return jsonify({
            "success":  True,
            "pdf_b64":  pdf_b64,
            "docx_b64": docx_b64
        })

    except Exception as e:
        import traceback
        print(f"An error occurred: {e}")
        print(traceback.format_exc())
        return jsonify({"error": f"Python service error: {str(e)}"}), 500
    finally:
        if original_mplconfigdir is not None:
            os.environ['MPLCONFIGDIR'] = original_mplconfigdir
        elif 'MPLCONFIGDIR' in os.environ:
            del os.environ['MPLCONFIGDIR']
        
        # shutil.rmtree(temp_dir, ignore_errors=True)


if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5001, debug=True)
